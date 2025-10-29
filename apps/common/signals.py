from django.db.models.signals import pre_save, post_save
from django.dispatch import receiver
from django.contrib.contenttypes.models import ContentType
from apps.common.models import AuditTrail
from apps.users.middleware import AuditUserMiddleware
from django.utils.functional import SimpleLazyObject
from apps.file_manager.models import File, FileChunk
from apps.tables.utils import get_model
from sentence_transformers import SentenceTransformer
from loader.models import is_pgvector_enabled

LOCAL_APPS = ['common']

@receiver(pre_save)
def log_model_changes(sender, instance, **kwargs):
    if sender == AuditTrail:
        return

    app_label = sender._meta.app_label
    if app_label not in LOCAL_APPS:
        return
    
    current_user = AuditUserMiddleware.get_current_user()
    if isinstance(current_user, SimpleLazyObject) and current_user.is_anonymous:
        current_user = None

    if instance.pk:
        old_instance = sender.objects.get(pk=instance.pk)

        for field in instance._meta.fields:
            field_name = field.name
            old_value = getattr(old_instance, field_name)
            new_value = getattr(instance, field_name)

            if old_value != new_value:
                audit_trail = AuditTrail.objects.create(
                    content_type=ContentType.objects.get_for_model(instance),
                    object_id=instance.pk,
                    field_name=field_name,
                    old_value=str(old_value),
                    new_value=str(new_value),
                )
                if current_user:
                    audit_trail.changed_by = current_user
                    audit_trail.save()


import os
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
from celery import shared_task
from django.contrib.postgres.search import SearchVector

def to_sql_vector(text: str, model: SentenceTransformer = None):
    if model is None:
        model = get_model()
    vec = model.encode([text], normalize_embeddings=True)[0].astype(float)
    return vec.tolist()


def split_with_position(text, chunk_size=2000, page_number=None):
    chunks = []
    total_len = len(text)
    for i, start in enumerate(range(0, total_len, chunk_size)):
        chunk_text = text[start:start+chunk_size]

        center = start + len(chunk_text) / 2
        rel = center / total_len
        if rel < 0.33:
            position = "top"
        elif rel < 0.66:
            position = "middle"
        else:
            position = "bottom"

        chunks.append({
            "text": chunk_text,
            "index": i,
            "page_number": page_number,
            "position": position,
        })
    return chunks


def extract_text(file_field):
    name = file_field.name.lower()
    ext = os.path.splitext(name)[1]

    try:
        if ext == ".txt":
            with file_field.open("r", encoding="utf-8") as f:
                return f.read()

        elif ext == ".pdf":
            with file_field.open("rb") as f:
                with pdfplumber.open(f) as pdf:
                    pages = []
                    for i, page in enumerate(pdf.pages, start=1):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            pages.append((i, page_text))
            return pages

        elif ext == ".docx":
            with file_field.open("rb") as f:
                doc = Document(f)
                full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return [(1, full_text)]

        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
            with file_field.open("rb") as f:
                img = Image.open(f)
                return pytesseract.image_to_string(img)

        else:
            return ""
    except Exception as e:
        print(f"Error extracting text from {name}: {e}")
        return ""


@shared_task
def create_file_chunks_task(file_id):
    try:
        file_instance = File.objects.get(pk=file_id)
    except File.DoesNotExist:
        print(f"File with ID {file_id} does not exist.")
        return

    content = extract_text(file_instance.file)
    if not content:
        print(f"No content extracted from file: {file_instance.file.name}")
        return
    
    chunk_size = 2000
    if isinstance(content, list): 
        for page_num, page_text in content:
            chunks = split_with_position(page_text, chunk_size, page_number=page_num)
            for ch in chunks:
                vector_value = to_sql_vector(ch["text"]) if is_pgvector_enabled() else ch["text"]
                chunk = FileChunk.objects.create(
                    file=file_instance,
                    vector=vector_value,
                    chunk_text=ch["text"],
                    chunk_index=ch["index"],
                    page_number=ch["page_number"],
                    position=ch["position"]
                )

                FileChunk.objects.filter(pk=chunk.pk).update(
                    fts=SearchVector("chunk_text")
                )
    else:
        chunks = split_with_position(content, chunk_size)
        for ch in chunks:
            vector_value = to_sql_vector(ch["text"]) if is_pgvector_enabled() else ch["text"]
            chunk = FileChunk.objects.create(
                file=file_instance,
                vector=vector_value,
                chunk_text=ch["text"],
                chunk_index=ch["index"],
                page_number=None,
                position=ch["position"]
            )

            FileChunk.objects.filter(pk=chunk.pk).update(
                fts=SearchVector("chunk_text")
            )

@receiver(post_save, sender=File)
def create_file_chunks(sender, instance, created, **kwargs):
    if created and instance.file:
        create_file_chunks_task.delay(instance.id)