import os
import json
import textwrap
import base64
from django.shortcuts import redirect, get_object_or_404, render
from apps.tables.models import (
    Tab, BaseCharts, ModelChoices,
    ChartPrompt, ChartType2, ScoreCard,
    RiskAssessment, BusinessImpactItem,
    OwnerRole
)
from datetime import datetime, timedelta
from django.utils import timezone
from django.contrib.contenttypes.models import ContentType
from home.models import IPE, ColumnOrder
from apps.common.models import SavedFilter, FieldType
from django.urls import reverse
from apps.tables.choices import *
from apps.tables.utils import (
    same_key_filter, common_date_filter,
    common_float_filter, common_integer_filter, common_count_filter, 
    common_unique_filter
)
from django.utils.html import strip_tags
from django.http import HttpResponse, JsonResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from io import BytesIO
from apps.graic.utils import get_agent, api_engine
from apps.file_manager.models import File, ActionStatus, FileManager
from django.core.files.base import ContentFile
from django.conf import settings
from apps.tables.models import DocumentStatus
from loader.models import InstantUpload

def create_risk_chart_filter(request, chart_id):
    chart = BaseCharts.objects.get(id=chart_id)
    content_type = ContentType.objects.get(app_label=chart.content_type.app_label, model=chart.content_type.model)
    model = content_type.model_class()

    field_type_map = {
        'date': getattr(model, 'date_fields_to_convert', []),
        'int': getattr(model, 'integer_fields', []),
        'float': getattr(model, 'float_fields', [])
    }

    def get_field_type(value):
        for f_type, fields in field_type_map.items():
            if value in fields:
                return getattr(FieldType, f_type.upper())
        return FieldType.TEXT

    def get_boolean_field(name, index):
        return request.POST.get(f'{name}_{index}') == 'true'
    
    ipe_path = reverse('risk_chart_details', args=[chart.id])
    last_ipe = IPE.objects.filter(path=ipe_path).last()
    base_description = last_ipe.description if last_ipe else ""

    existing_count_filter = SavedFilter.objects.filter(
        userID=request.user.id, parent=ModelChoices.CHART, is_count=True, chart_id=chart.id
    ).first()

    filters_description = []

    for key, value in request.POST.items():
        if key.startswith('field_name_') and value:
            index = key.split('_')[-1]
            field_name = value
            field_type = get_field_type(value)

            if field_type == FieldType.TEXT:
                value_start = request.POST.get(f'value_start_{index}') or ''
                value_end = request.POST.get(f'value_end_{index}') or ''
            else:
                value_start = request.POST.get(f'value_start_{index}') or None
                value_end = request.POST.get(f'value_end_{index}') or None

            is_null = get_boolean_field('is_null', index)
            is_not_null = get_boolean_field('is_not_null', index)
            is_not = get_boolean_field('is_not', index)
            is_unique = get_boolean_field('is_unique', index)
            is_count = get_boolean_field('is_count', index)

            filter_id = request.POST.get(f'filter_id_{index}')
            if filter_id in [None, '', 'null']:
                filter_id = None
            
            if is_count:
                if existing_count_filter and (not filter_id or str(existing_count_filter.id) != filter_id):
                    is_count = False

            filter_data = {
                'userID': request.user.id,
                'parent': ModelChoices.CHART,
                'field_name': field_name,
                'field_type': field_type,
                'value_start': value_start,
                'value_end': value_end,
                'is_null': is_null,
                'is_not_null': is_not_null,
                'is_not': is_not,
                'is_unique': is_unique,
                'is_count': is_count,
                'chart_id': chart.id
            }

            if filter_id:
                SavedFilter.objects.filter(pk=filter_id).update(**filter_data)
                saved_filter = SavedFilter.objects.get(pk=filter_id)
            else:
                saved_filter = SavedFilter.objects.create(**filter_data)
            
            if not chart.saved_filters.filter(pk=saved_filter.pk).exists():
                chart.saved_filters.add(saved_filter)
            
            filters_description.append(
                f"Field: {field_name}, Start: {value_start}, End: {value_end}, "
                f"Null: {is_null}, Not Null: {is_not_null}, Not: {saved_filter.is_not}, Unique: {is_unique}, Count: {is_count}"
            )
    
    user = request.user
    timestamp_utc = timezone.now()
    timestamp_local = timezone.localtime(timestamp_utc)
    timestamp_str = timestamp_local.strftime('%Y-%m-%d %H:%M:%S')

    new_description = (
        f"{base_description}\n" if base_description else ""
    ) + f'Filters applied by: {user.get_full_name() or user.username}\n' \
        f'Applied on: {timestamp_str}\n' \
        f'Filters:\n' + "\n".join(filters_description)

    IPE.objects.create(
        userID=user,
        path=ipe_path,
        chart_id=chart.id,
        description=new_description
    )
    
    return redirect(request.META.get('HTTP_REFERER'))

def add_risk_card(request, tab_id):
    tab = get_object_or_404(Tab, id=tab_id)

    if request.method == 'POST':
        base_view = request.POST.get('base_view')
        form_data = {}

        for key, value in request.POST.items():
            if key in ('csrfmiddlewaretoken', 'business_impacts', 'base_view', 'owner_roles'):
                continue
            form_data[key] = value

        base_chart, created = BaseCharts.objects.get_or_create(
            base_view=base_view,
            chart_type=ChartType2.RISK_CHART,
            defaults={'content_type': tab.content_type, 'name': 'Score Cards'}
        )

        form_data['base_chart'] = base_chart

        risk_assessment, created = RiskAssessment.objects.update_or_create(
            parent_tab=tab,
            defaults=form_data
        )
        ScoreCard.objects.get_or_create(risk_card=risk_assessment)

        business_impacts = BusinessImpactItem.objects.filter(
            code__in=request.POST.getlist('business_impacts')
        )
        risk_assessment.business_impacts.set(business_impacts)

        owner_roles = OwnerRole.objects.filter(
            name__in=request.POST.getlist('owner_roles')
        )
        risk_assessment.recommended_owner_role.set(owner_roles)

        return redirect(request.META.get('HTTP_REFERER'))

    return redirect(request.META.get('HTTP_REFERER'))

from collections import defaultdict
from apps.tables.choices import RISK_MATRIX, CONFIDENCE_PERCENTAGE

def normalize_enum(value):
    if not value:
        return ""

    if "-" in value:
        value = value.split("-")[-1]

    return value.strip().upper()

def risk_chart_details(request, chart_id):
    base_chart = get_object_or_404(BaseCharts, id=chart_id)
    chart_prompt, _ = ChartPrompt.objects.get_or_create(base_chart=base_chart)

    tabs = Tab.objects.filter(
        base_view=base_chart.base_view
    ).order_by("created_at")

    tab_heatmaps = []

    for tab in tabs:
        scorecards = (
            ScoreCard.objects.filter(
                risk_card__base_chart=base_chart,
                risk_card__parent_tab=tab
            )
            .select_related("risk_card")
        )

        if not scorecards.exists():
            continue

        heatmap_data = defaultdict(list)

        for scorecard in scorecards:
            risk = scorecard.risk_card

            likelihood = normalize_enum(risk.likelihood)
            residual = normalize_enum(risk.residual_risk)

            key = f"{likelihood}|{residual}"

            heatmap_data[key].append({
                "id": str(risk.id),
                "name": risk.name,
                "confidence": risk.confidence,
                "confidence_value": CONFIDENCE_PERCENTAGE.get(
                    risk.confidence, 0
                ),
            })

        cell_confidence = {}
        for key, risks in heatmap_data.items():
            total = sum(r["confidence_value"] for r in risks)
            cell_confidence[key] = round(total / len(risks)) if risks else 0

        tab_heatmaps.append({
            "tab": tab,
            "risk": scorecards.first().risk_card,
            "score_card": scorecards.first(),
            "risk_matrix": RISK_MATRIX,
            "heatmap_data": dict(heatmap_data),
            "cell_confidence": cell_confidence,
        })

    context = {
        "base_chart": base_chart,
        "tabs": tabs,
        "charts": BaseCharts.objects.filter(base_view=base_chart.base_view, chart_type=ChartType2.BASE_CHART),
        "risk_charts": BaseCharts.objects.filter(base_view=base_chart.base_view, chart_type=ChartType2.RISK_CHART),
        "chart_prompt": chart_prompt,
        "tab_heatmaps": tab_heatmaps,

        "inherent_impact": InherentImpact.choices,
        "likelihood": Likelihood.choices,
        "residual_risk": ResidualRiskRating.choices,
        "confidence": ConfidenceInResults.choices,
        "primary_root_cause": PrimaryRootCause.choices,
        "secondary_root_cause": SecondaryRootCause.choices,
        "business_impact": BusinessImpact.choices,
        "recommended_owner_role": OwnerRoleChoices.choices,
    }

    return render(request, "apps/charts/risk_chart_details.html", context)

def delete_score_card(request, score_card_id):
    score_card = get_object_or_404(ScoreCard, id=score_card_id)
    score_card.delete()
    return redirect(request.META.get('HTTP_REFERER'))


# Export

from types import SimpleNamespace
from django.http import QueryDict

def fake_request_for_user(user):
    req = SimpleNamespace()
    req.user = user
    req.GET = QueryDict("")
    return req

def decode_base64_image(data):
    if not data:
        return None

    header, encoded = data.split(",", 1)
    return BytesIO(base64.b64decode(encoded))

def quill_delta_to_text(value):
    if not value:
        return ""

    try:
        value = strip_tags(value or "").strip()
        if isinstance(value, dict):
            delta = value
        else:
            delta = json.loads(value)

        text = []
        for op in delta.get("ops", []):
            insert = op.get("insert")
            if isinstance(insert, str):
                text.append(insert)

        return "".join(text).strip()
    except Exception:
        return ""

def serialize_risk_assessment(risk: RiskAssessment):
    data = {
        "Risk Name": risk.name,
        "Risk Description": quill_delta_to_text(risk.description),
        "Inherent Impact": risk.get_inherent_impact_display(),
        "Likelihood": risk.get_likelihood_display(),
        "Residual Risk": risk.get_residual_risk_display(),
        "Confidence": risk.get_confidence_display(),
        "Primary Root Cause": risk.get_primary_root_cause_display(),
        "Secondary Root Cause": (
            risk.get_secondary_root_cause_display()
            if risk.secondary_root_cause else "N/A"
        ),
        "Business Impacts": ", ".join(
            bi.code for bi in risk.business_impacts.all()
        ) or "N/A",
        "Audit Recommendation": quill_delta_to_text(risk.audit_recommendation),
        "Recommended Owner Roles": ", ".join(
            role.name for role in risk.recommended_owner_role.all()
        ) or "N/A",
        "Target Remediation Date": risk.target_remediation_date.strftime("%Y-%m-%d"),
        "Management Response": risk.management_response or "N/A",
        "Agreed Action Plan": risk.agreed_action_plan or "N/A",
        "Created At": risk.created_at.strftime("%Y-%m-%d %H:%M:%S"),
        "Updated At": risk.updated_at.strftime("%Y-%m-%d %H:%M:%S"),
    }
    return data

def build_risk_prompt(risk_data, custom_prompt, agent_data):
    current_dt = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    system_prompt = agent_data["system_prompt"]
    if custom_prompt:
        system_prompt = f"{system_prompt}\n\n{custom_prompt}"

    fields_text = "\n".join(
        f"{key}: {value}" for key, value in risk_data.items()
    )

    parts = [
        current_dt,
        system_prompt,
        agent_data.get("instruction", ""),
        "Risk Assessment Details:",
        fields_text,
    ]

    return "\n\n".join(parts)

def get_or_create_folder(name, parent=None, user=None):
    folder, _ = FileManager.objects.get_or_create(
        name=name,
        parent=parent,
        defaults={
            "created_by": user,
            "updated_by": user,
            "action_status": ActionStatus.IS_ACTIVE,
            "folder_status": DocumentStatus.APPROVED,
        }
    )
    return folder


def get_chart_save_path(user, buffer, filename):
    buffer.seek(0)
    username = user.username if user.is_authenticated else "anonymous"

    users_root = get_or_create_folder("users", parent=None, user=user)
    user_folder = get_or_create_folder(username, parent=users_root, user=user)
    chart_folder = get_or_create_folder("risk", parent=user_folder, user=user)

    file_obj = File.objects.create(
        file_manager=chart_folder,
        file=ContentFile(buffer.getvalue(), name=filename),
        action_status=ActionStatus.IS_ACTIVE,
        file_status=DocumentStatus.APPROVED,
        uploaded_by=user,
        updated_by=user
    )

    base_path = os.path.join(settings.MEDIA_ROOT, "users", username, "risk")
    os.makedirs(base_path, exist_ok=True)
    file_path = os.path.join(base_path, filename)
    with open(file_path, "wb") as f:
        f.write(buffer.getvalue())

    return {
        "file_id": file_obj.id,
        "filename": filename,
        "path": f"users/{username}/risk"
    }

def get_tab_queryset_for_ai(request, tab: Tab):
    content_type = ContentType.objects.get(
        app_label=tab.content_type.app_label,
        model=tab.content_type.model
    )
    tab_model = content_type.model_class()

    db_field_names = [
        field.name for field in tab_model._meta.get_fields()
        if not field.is_relation
    ]

    try:
        user_order = ColumnOrder.objects.get(
            user=request.user,
            table_name=f'{tab_model.__name__}',
            tab_id=str(tab.id)
        )
        ordered_fields = [
            col['key'] for col in user_order.column_order
            if col['key'] is not None
        ]
    except ColumnOrder.DoesNotExist:
        ordered_fields = db_field_names

    combined_q_objects, user_unique_filter, user_query_conditions, user_count_filters = same_key_filter(tab.saved_filters.filter(field_type=FieldType.TEXT), return_count_filters=True)

    filter_string = {}

    date_string, date_unique_filter, date_query_conditions, date_count_filters = common_date_filter(tab.saved_filters.filter(field_type=FieldType.DATE), return_count_filters=True)
    filter_string.update(date_string)

    int_string, int_unique_filter, int_query_conditions, int_count_filters = common_integer_filter(tab.saved_filters.filter(field_type=FieldType.INT), return_count_filters=True)
    filter_string.update(int_string)

    float_string, float_unique_filter, float_query_conditions, float_count_filters = common_float_filter(tab.saved_filters.filter(field_type=FieldType.FLOAT), return_count_filters=True)
    filter_string.update(float_string)

    snapshot_filter = {}

    try:
        content_type = ContentType.objects.get(model=tab_model.__name__.lower())
        snapshots = InstantUpload.objects.filter(content_type=content_type).order_by('-created_at')
        snapshot = request.GET.get('snapshot')

        if snapshot and not snapshot == 'all':
            summary = InstantUpload.objects.get(id=snapshot)
            snapshot_filter['loader_instance'] = summary.pk
        
        elif snapshot and snapshot == 'all':
            snapshot_filter= {}
        
        elif tab.snapshot == 'latest':
            latest_snapshot = snapshots.latest('created_at')
            snapshot_filter['loader_instance'] = latest_snapshot.pk
        
        elif tab.snapshot and not tab.snapshot == 'all':
            summary = InstantUpload.objects.get(id=int(tab.snapshot))
            snapshot_filter['loader_instance'] = int(tab.snapshot)
        
        elif tab.snapshot and tab.snapshot == 'all':
            snapshot_filter= {}
        
    except:
        pass

    queryset = (
        tab_model.objects
        .filter(combined_q_objects)
        .filter(**filter_string)
        .filter(**snapshot_filter)
    )

    base_queryset = tab_model.objects.filter(combined_q_objects).filter(**filter_string).filter(**snapshot_filter)
    queryset = base_queryset
    if hasattr(tab_model, 'parent'):
        queryset = queryset.filter(parent=None)
        try:
            queryset = queryset.filter(action_status=ActionStatus.IS_ACTIVE)
        except:
            pass

    if user_query_conditions:
        queryset = queryset.filter(user_query_conditions)
    if date_query_conditions:
        queryset = queryset.filter(date_query_conditions)
    if int_query_conditions:
        queryset = queryset.filter(int_query_conditions)
    if float_query_conditions:
        queryset = queryset.filter(float_query_conditions)
    
    if user_count_filters:
        queryset = common_count_filter(user_count_filters, base_queryset, queryset, ordered_fields)
    elif date_count_filters:
        queryset = common_count_filter(date_count_filters, base_queryset, queryset, ordered_fields)
    elif int_count_filters:
        queryset = common_count_filter(int_count_filters, base_queryset, queryset, ordered_fields)
    elif float_count_filters:
        queryset = common_count_filter(float_count_filters, base_queryset, queryset, ordered_fields)
    else:
        if order_by in ['count', '-count']:
            order_by = 'pk'

    table_name = f"{tab.content_type.app_label}_{tab.content_type.model}"
    if user_unique_filter:
        queryset = common_unique_filter(request, user_unique_filter, queryset, snapshot_filter, table_name)
    if date_unique_filter:
        queryset = common_unique_filter(request, date_unique_filter, queryset, snapshot_filter, table_name)
    if int_unique_filter:
        queryset = common_unique_filter(request, int_unique_filter, queryset, snapshot_filter, table_name)
    if float_unique_filter:
        queryset = common_unique_filter(request, float_unique_filter, queryset, snapshot_filter, table_name)

    order_by = tab.order_by or "pk"
    if "similarity" in order_by:
        order_by = "pk"

    queryset = queryset.order_by(order_by)

    return queryset, ordered_fields

def serialize_queryset_for_ai(queryset, fields):
    rows = []
    for obj in queryset:
        row = []
        for field in fields:
            val = getattr(obj, field, "")
            if isinstance(val, datetime):
                val = val.strftime("%Y-%m-%d %H:%M:%S")
            row.append(f"{field}={val}")
        rows.append(" | ".join(row))
    return rows

def chunk_list(data, chunk_size=50):
    for i in range(0, len(data), chunk_size):
        yield data[i:i + chunk_size]


def generate_ai_risk_analysis(request, risk: RiskAssessment, custom_prompt: str = "", image_data=None):
    tab = risk.parent_tab
    agent_data = get_agent("risk_agent")

    queryset, ordered_fields = get_tab_queryset_for_ai(request, tab)
    rows = serialize_queryset_for_ai(queryset, ordered_fields)

    ai_chunk_outputs = []

    for idx, chunk in enumerate(chunk_list(rows, 50)):
        dataset_prompt = f"Dataset chunk {idx + 1}\n\n" + "\n".join(chunk)
        full_prompt = build_risk_prompt(serialize_risk_assessment(risk), custom_prompt, agent_data) + "\n\n" + dataset_prompt

        chunk_response = api_engine(
            prompt=full_prompt,
            is_online=True,
            llm=agent_data.get("llm", "openai"),
            image_data=image_data
        )
        ai_chunk_outputs.append(chunk_response)

    if len(ai_chunk_outputs) == 0:
        return ""
    elif len(ai_chunk_outputs) == 1:
        return ai_chunk_outputs[0]

    consolidation_prompt = (
        "You are given multiple partial risk analyses based on dataset chunks.\n"
        "Merge them into ONE clear, structured, professional risk assessment.\n\n"
        + "\n\n".join(ai_chunk_outputs)
    )

    final_analysis = api_engine(
        prompt=consolidation_prompt,
        is_online=True,
        llm=agent_data.get("llm", "openai"),
        image_data=image_data
    )

    return final_analysis


def export_risk_docx(request):
    if request.method != "POST":
        return HttpResponse("Invalid request", status=400)

    heatmap_image_base64 = request.POST.get("heatmap_image")
    heatmap_image_file = decode_base64_image(heatmap_image_base64)
    risk_id = request.POST.get("risk_id")
    custom_prompt = request.POST.get("custom_prompt", "").strip()
    export_option = request.POST.get("export_option", "download")

    risk = get_object_or_404(RiskAssessment, id=risk_id)
    ai_analysis = generate_ai_risk_analysis(request, risk, custom_prompt, image_data=heatmap_image_base64)
    risk_data = serialize_risk_assessment(risk)

    doc = Document()
    doc.add_heading(f"Risk Assessment Report – {risk.name}", level=1)

    if heatmap_image_file:
        from docx.shared import Inches

        doc.add_heading("Risk Heatmap", level=2)
        doc.add_picture(heatmap_image_file, width=Inches(6))

    doc.add_heading("Risk Details", level=2)
    for key, value in risk_data.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("grAIc Analysis", level=2)
    doc.add_paragraph(ai_analysis or "No analysis generated.")

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{risk.name}_{timestamp}.docx"

    if export_option == "grc":
        saved_file = get_chart_save_path(request.user, output, filename)
        return JsonResponse({"status": "saved", **saved_file})

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def export_risk_pdf(request):
    if request.method != "POST":
        return HttpResponse("Invalid request", status=400)

    heatmap_image_base64 = request.POST.get("heatmap_image")
    heatmap_image_file = decode_base64_image(heatmap_image_base64)
    risk_id = request.POST.get("risk_id")
    custom_prompt = request.POST.get("custom_prompt", "").strip()
    export_option = request.POST.get("export_option", "download")

    risk = get_object_or_404(RiskAssessment, id=risk_id)
    ai_analysis = generate_ai_risk_analysis(request, risk, custom_prompt, image_data=heatmap_image_base64)
    risk_data = serialize_risk_assessment(risk)

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    y = height - 40
    p.setFont("Helvetica-Bold", 18)
    p.drawString(40, y, f"Risk Assessment Report – {risk.name}")
    y -= 30

    if heatmap_image_file:
        from reportlab.lib.utils import ImageReader

        img = ImageReader(heatmap_image_file)
        iw, ih = img.getSize()
        aspect = ih / float(iw)

        img_width = width - 80
        img_height = img_width * aspect

        if y - img_height < 60:
            p.showPage()
            y = height - 40

        p.drawImage(img, 40, y - img_height, img_width, img_height)
        y -= img_height + 20

    y -= 10
    p.setFont("Helvetica-Bold", 14)
    p.drawString(40, y, "Risk Details")
    y -= 20

    p.setFont("Helvetica", 11)
    for key, value in risk_data.items():
        for line in textwrap.wrap(f"{key}: {value}", 95):
            p.drawString(40, y, line)
            y -= 14
            if y < 60:
                p.showPage()
                y = height - 40
                p.setFont("Helvetica", 11)

    y -= 10
    p.setFont("Helvetica-Bold", 14)
    p.drawString(40, y, "grAIc Analysis")
    y -= 20

    p.setFont("Helvetica", 11)
    for line in (ai_analysis or "").split("\n"):
        for wrapped in textwrap.wrap(line, 95):
            p.drawString(40, y, wrapped)
            y -= 14
            if y < 60:
                p.showPage()
                y = height - 40
                p.setFont("Helvetica", 11)

    p.showPage()
    p.save()

    pdf = buffer.getvalue()
    buffer.close()

    timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{risk.name}_{timestamp}.pdf"

    if export_option == "grc":
        saved_file = get_chart_save_path(request.user, BytesIO(pdf), filename)
        return JsonResponse({"status": "saved", **saved_file})

    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# Schedule export
from django.views.decorators.http import require_POST
from django.utils.dateparse import parse_datetime
from apps.tables.models import ScheduledRiskExport
from celery import shared_task

@require_POST
def create_risk_export_schedule(request):
    risk = get_object_or_404(RiskAssessment, id=request.POST.get("risk_id"))
    weekdays = request.POST.getlist("weekdays[]")
    chart_image = request.POST.get("chart_image", "")

    schedule = ScheduledRiskExport.objects.create(
        user=request.user,
        risk=risk,
        chart_image=chart_image,
        export_type=request.POST.get("export_type"),
        export_option="grc",
        frequency=request.POST.get("frequency"),
        start_at=parse_datetime(request.POST.get("start_at")),
        end_at=parse_datetime(request.POST.get("end_at")) if request.POST.get("end_at") else None,
        hour_interval=request.POST.get("hour_interval") or None,
        time_of_day=request.POST.get("time_of_day") or None,
        month_day=request.POST.get("month_day") or None,
        weekdays=weekdays,
        custom_prompt=request.POST.get("custom_prompt", ""),
    )

    return JsonResponse({
        "status": "scheduled",
        "schedule_id": schedule.id
    })


def generate_risk_export_from_schedule(schedule: ScheduledRiskExport):
    risk = schedule.risk
    user = schedule.user
    image_data = schedule.chart_image or None
    custom_prompt = schedule.custom_prompt or ""
    export_type = schedule.export_type
    export_option = schedule.export_option

    if export_type == "docx":
        output = export_risk_docx_programmatic(risk, user, custom_prompt, image_data, export_option)
    else:
        output = export_risk_pdf_programmatic(risk, user, custom_prompt, image_data, export_option)

    return output

def export_risk_docx_programmatic(risk, user, custom_prompt="", image_data=None, export_option="grc"):
    request = fake_request_for_user(user)
    heatmap_image_file = decode_base64_image(image_data)
    ai_analysis = generate_ai_risk_analysis(request=request, risk=risk, custom_prompt=custom_prompt, image_data=image_data)
    risk_data = serialize_risk_assessment(risk)

    doc = Document()
    doc.add_heading(f"Risk Assessment Report – {risk.name}", level=1)

    if heatmap_image_file:
        from docx.shared import Inches
        doc.add_heading("Risk Heatmap", level=2)
        doc.add_picture(heatmap_image_file, width=Inches(6))

    doc.add_heading("Risk Details", level=2)
    for key, value in risk_data.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("grAIc Analysis", level=2)
    doc.add_paragraph(ai_analysis or "No analysis generated.")

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")
    filename = f"risk_{risk.id}_{timestamp}.docx"

    if export_option == "grc":
        return get_chart_save_path(user, output, filename)

    return output.getvalue(), filename


def export_risk_pdf_programmatic(risk, user, custom_prompt="", image_data=None, export_option="grc"):
    request = fake_request_for_user(user)
    heatmap_image_file = decode_base64_image(image_data)
    ai_analysis = generate_ai_risk_analysis(request=request, risk=risk, custom_prompt=custom_prompt, image_data=image_data)
    risk_data = serialize_risk_assessment(risk)

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    y = height - 40
    p.setFont("Helvetica-Bold", 18)
    p.drawString(40, y, f"Risk Assessment Report – {risk.name}")
    y -= 30

    if heatmap_image_file:
        from reportlab.lib.utils import ImageReader
        img = ImageReader(heatmap_image_file)
        iw, ih = img.getSize()
        aspect = ih / float(iw)
        img_width = width - 80
        img_height = img_width * aspect
        if y - img_height < 60:
            p.showPage()
            y = height - 40
        p.drawImage(img, 40, y - img_height, img_width, img_height)
        y -= img_height + 20

    y -= 10
    p.setFont("Helvetica-Bold", 14)
    p.drawString(40, y, "Risk Details")
    y -= 20

    p.setFont("Helvetica", 11)
    for key, value in risk_data.items():
        for line in textwrap.wrap(f"{key}: {value}", 95):
            p.drawString(40, y, line)
            y -= 14
            if y < 60:
                p.showPage()
                y = height - 40
                p.setFont("Helvetica", 11)

    y -= 10
    p.setFont("Helvetica-Bold", 14)
    p.drawString(40, y, "grAIc Analysis")
    y -= 20

    p.setFont("Helvetica", 11)
    for line in (ai_analysis or "").split("\n"):
        for wrapped in textwrap.wrap(line, 95):
            p.drawString(40, y, wrapped)
            y -= 14
            if y < 60:
                p.showPage()
                y = height - 40
                p.setFont("Helvetica", 11)

    p.showPage()
    p.save()

    pdf = buffer.getvalue()
    buffer.close()

    timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")
    filename = f"risk_{risk.id}_{timestamp}.pdf"

    if export_option == "grc":
        return get_chart_save_path(user, BytesIO(pdf), filename)

    return pdf, filename

@shared_task
def run_scheduled_risk_exports():
    now = timezone.now()
    schedules = ScheduledRiskExport.objects.filter(
        is_active=True,
        start_at__lte=now,
    )

    for schedule in schedules:
        if schedule.end_at and now > schedule.end_at:
            schedule.is_active = False
            schedule.save(update_fields=["is_active"])
            continue

        if not should_run_now_risk(schedule, now):
            continue

        generate_risk_export_from_schedule(schedule)

        schedule.last_run_at = now
        if schedule.frequency == "once":
            schedule.is_active = False

        schedule.save(update_fields=["last_run_at", "is_active"])

def should_run_now_risk(schedule, now):
    if schedule.frequency == "once":
        return schedule.last_run_at is None

    if schedule.frequency == "hourly":
        if not schedule.last_run_at:
            return True
        return (now - schedule.last_run_at).total_seconds() >= (schedule.hour_interval or 1) * 3600

    if schedule.frequency == "daily":
        return now.hour == (schedule.time_of_day.hour if schedule.time_of_day else now.hour) and (
            schedule.last_run_at is None or schedule.last_run_at.date() < now.date()
        )

    if schedule.frequency == "weekly":
        return now.weekday() in schedule.weekdays and now.hour == (schedule.time_of_day.hour if schedule.time_of_day else now.hour) and (
            not schedule.last_run_at or schedule.last_run_at.date() < now.date()
        )

    if schedule.frequency == "monthly":
        return now.day == (schedule.month_day or now.day) and now.hour == (schedule.time_of_day.hour if schedule.time_of_day else now.hour) and (
            not schedule.last_run_at or schedule.last_run_at.month < now.month
        )

    return False
