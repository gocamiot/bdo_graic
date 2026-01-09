import os
import json
from django.core.serializers.json import DjangoJSONEncoder
from django.shortcuts import redirect, get_object_or_404, render
from apps.tables.models import (
    Tab, BaseCharts, TabCharts,
    SelectedRows, ModelChoices,
    ActionStatus, DocumentStatus
)
from apps.tables.utils import (
    software_filter, same_key_filter, common_date_filter,
    common_float_filter, common_integer_filter, common_count_filter, 
    common_unique_filter, get_model_fields, get_user_id
)
from django.utils import timezone
from django.contrib.contenttypes.models import ContentType
from home.models import ColumnOrder, IPE
from apps.common.models import SavedFilter, FieldType
from loader.models import InstantUpload
from dateutil.parser import isoparse
from django.db.models import F, Case, When, IntegerField
from datetime import datetime
from collections import Counter
from django.urls import reverse
from django.conf import settings
from django.http import JsonResponse
from apps.file_manager.models import File, ActionStatus, FileManager
from django.core.files.base import ContentFile

def create_chart_filter(request, chart_id):
    chart = BaseCharts.objects.get(id=chart_id)
    content_type = ContentType.objects.get(app_label=chart.content_type.app_label, model=chart.content_type.model)
    model = content_type.model_class()

    field_type_map = {
        'date': getattr(model, 'date_fields_to_convert', []),
        'int': getattr(model, 'integer_fields', []),
        'float': getattr(model, 'float_fields', [])
    }

    def get_field_type(value):
        for f_type, fields in field_type_map.items():
            if value in fields:
                return getattr(FieldType, f_type.upper())
        return FieldType.TEXT

    def get_boolean_field(name, index):
        return request.POST.get(f'{name}_{index}') == 'true'
    
    ipe_path = reverse('chart_details', args=[chart.id])
    last_ipe = IPE.objects.filter(path=ipe_path).last()
    base_description = last_ipe.description if last_ipe else ""

    existing_count_filter = SavedFilter.objects.filter(
        userID=request.user.id, parent=ModelChoices.CHART, is_count=True, chart_id=chart.id
    ).first()

    filters_description = []

    for key, value in request.POST.items():
        if key.startswith('field_name_') and value:
            index = key.split('_')[-1]
            field_name = value
            field_type = get_field_type(value)

            if field_type == FieldType.TEXT:
                value_start = request.POST.get(f'value_start_{index}') or ''
                value_end = request.POST.get(f'value_end_{index}') or ''
            else:
                value_start = request.POST.get(f'value_start_{index}') or None
                value_end = request.POST.get(f'value_end_{index}') or None

            is_null = get_boolean_field('is_null', index)
            is_not_null = get_boolean_field('is_not_null', index)
            is_not = get_boolean_field('is_not', index)
            is_unique = get_boolean_field('is_unique', index)
            is_count = get_boolean_field('is_count', index)

            filter_id = request.POST.get(f'filter_id_{index}')
            if filter_id in [None, '', 'null']:
                filter_id = None
            
            if is_count:
                if existing_count_filter and (not filter_id or str(existing_count_filter.id) != filter_id):
                    is_count = False

            filter_data = {
                'userID': request.user.id,
                'parent': ModelChoices.CHART,
                'field_name': field_name,
                'field_type': field_type,
                'value_start': value_start,
                'value_end': value_end,
                'is_null': is_null,
                'is_not_null': is_not_null,
                'is_not': is_not,
                'is_unique': is_unique,
                'is_count': is_count,
                'chart_id': chart.id
            }

            if filter_id:
                SavedFilter.objects.filter(pk=filter_id).update(**filter_data)
                saved_filter = SavedFilter.objects.get(pk=filter_id)
            else:
                saved_filter = SavedFilter.objects.create(**filter_data)
            
            if not chart.saved_filters.filter(pk=saved_filter.pk).exists():
                chart.saved_filters.add(saved_filter)
            
            filters_description.append(
                f"Field: {field_name}, Start: {value_start}, End: {value_end}, "
                f"Null: {is_null}, Not Null: {is_not_null}, Not: {saved_filter.is_not}, Unique: {is_unique}, Count: {is_count}"
            )
    
    user = request.user
    timestamp_utc = timezone.now()
    timestamp_local = timezone.localtime(timestamp_utc)
    timestamp_str = timestamp_local.strftime('%Y-%m-%d %H:%M:%S')

    new_description = (
        f"{base_description}\n" if base_description else ""
    ) + f'Filters applied by: {user.get_full_name() or user.username}\n' \
        f'Applied on: {timestamp_str}\n' \
        f'Filters:\n' + "\n".join(filters_description)

    IPE.objects.create(
        userID=user,
        path=ipe_path,
        chart_id=chart.id,
        description=new_description
    )
    
    return redirect(request.META.get('HTTP_REFERER'))

def add_chart(request, tab_id):
    tab = get_object_or_404(Tab, id=tab_id)
    if request.method == 'POST':
        base_view = request.POST.get('base_view')
        name = request.POST.get('name')
        x_field = request.POST.get('x_axis')
        y_field = request.POST.get('y_axis')
        chart_type = request.POST.get('chart_type')
        color = request.POST.get('color')

        base_chart, created = BaseCharts.objects.get_or_create(base_view=base_view)
        if created:
            base_chart.content_type = tab.content_type
            base_chart.save()

        TabCharts.objects.create(
            base_chart=base_chart,
            parent_tab=tab,
            name=name,
            x_field=x_field,
            y_field=y_field,
            chart_type=chart_type,
            color=color,
            created_at=timezone.now(),
            updated_at=timezone.now()
        )

        return redirect(request.META.get('HTTP_REFERER'))
    
    return redirect(request.META.get('HTTP_REFERER'))


def chart_details(request, chart_id):
    base_chart = get_object_or_404(BaseCharts, id=chart_id)
    chart_model = None
    pre_column = ('loader_instance', 'ID', 'json_data', 'fts', 'hash_data')

    base_saved_filters = list(SavedFilter.objects.filter(
        parent=ModelChoices.CHART, chart_id=base_chart.id
    ).values())
    for filter in base_saved_filters:
        if 'created_at' in filter:
            filter['created_at'] = filter['created_at'].isoformat()
    base_saved_filters_json = json.dumps(base_saved_filters)

    chart_data = {}
    chart_ids = []
    for chart in TabCharts.objects.filter(base_chart=base_chart).order_by('-created_at'):
        chart_ids.append(str(chart.id))

        tab = chart.parent_tab

        tab_saved_count_filter = SavedFilter.objects.filter(
            userID=get_user_id(request),
            parent=ModelChoices.TAB,
            tab_id=tab.id,
            is_count=True
        ).last()

        if tab_saved_count_filter and chart.y_field == 'count':
            chart.x_field = tab_saved_count_filter.field_name
            chart.save()

        content_type = ContentType.objects.get(app_label=tab.content_type.app_label, model=tab.content_type.model)
        db_field_names = [field.name for field in content_type.model_class()._meta.get_fields() if not field.is_relation]
        chart_model = content_type.model_class()

        try:
            user_order = ColumnOrder.objects.get(user=request.user, table_name=f'{chart_model.__name__}', tab_id=f"{tab.id}")
            column_names = [col['key'] for col in user_order.column_order if col['key'] is not None]
            ordered_fields = column_names
        except ColumnOrder.DoesNotExist:
            ordered_fields = db_field_names

        field_dict = {field.key: field for field in tab.hide_show_filters.all()}
        field_names = [field_dict[key] for key in ordered_fields if key in field_dict]

        selected_rows = SelectedRows.objects.filter(
            model=f'{tab.content_type.model}', 
            model_choice=ModelChoices.TAB,
            tab_id=tab.id
        ).values_list('rows', flat=True)
        selected_rows = [int(item) for row in selected_rows for item in row.split(',')]

        # Snapshot
        latest_snapshot = ''
        summary = None
        snapshot_filter = {}

        if not tab.is_dynamic_query:
            try:
                content_type = ContentType.objects.get(model=chart_model.__name__.lower())
                snapshots = InstantUpload.objects.filter(content_type=content_type).order_by('-created_at')
                snapshot = request.GET.get('snapshot')

                if snapshot and not snapshot == 'all':
                    summary = InstantUpload.objects.get(id=snapshot)
                    snapshot_filter['loader_instance'] = summary.pk
                
                elif snapshot and snapshot == 'all':
                    snapshot_filter= {}
                
                elif tab.snapshot == 'latest':
                    latest_snapshot = snapshots.latest('created_at')
                    snapshot_filter['loader_instance'] = latest_snapshot.pk
                
                elif tab.snapshot and not tab.snapshot == 'all':
                    summary = InstantUpload.objects.get(id=int(tab.snapshot))
                    snapshot_filter['loader_instance'] = int(tab.snapshot)
                
                elif tab.snapshot and tab.snapshot == 'all':
                    snapshot_filter= {}
                
            except:
                pass
        
        else:
            snapshots = chart_model.objects.exclude(snapshot=None).values('snapshot').distinct()

        filter_string = {}
        pre_filters = {}
        if tab.pre_filters:
            pre_filters = eval(tab.pre_filters)
        
        if tab.is_dynamic_query:
            query_snapshot = tab.query_snapshot
            if request.GET.get('query_snapshot'):
                query_snapshot = request.GET.get('query_snapshot')
            else:
                query_snapshot = tab.query_snapshot
            

            if query_snapshot and query_snapshot != 'all':
                parsed_datetime = isoparse(query_snapshot)
                snapshot_filter['snapshot'] = parsed_datetime
        
        # TEXT FILTERS
        tab_text_filters = tab.saved_filters.filter(field_type=FieldType.TEXT)
        chart_text_filters = SavedFilter.objects.filter(
            userID=get_user_id(request),
            parent=ModelChoices.CHART,
            chart_id=base_chart.id,
            field_type=FieldType.TEXT
        )
        merged_text_filters = tab_text_filters | chart_text_filters
        combined_q_objects, user_unique_filter, user_query_conditions, user_count_filters = same_key_filter(
            merged_text_filters, return_count_filters=True
        )

        # DATE FILTERS
        tab_date_filters = tab.saved_filters.filter(field_type=FieldType.DATE)
        chart_date_filters = SavedFilter.objects.filter(
            userID=get_user_id(request),
            parent=ModelChoices.CHART,
            chart_id=base_chart.id,
            field_type=FieldType.DATE
        )
        merged_date_filters = tab_date_filters | chart_date_filters
        date_string, date_unique_filter, date_query_conditions, date_count_filters = common_date_filter(
            merged_date_filters, return_count_filters=True
        )
        filter_string.update(date_string)

        # INT FILTERS
        tab_int_filters = tab.saved_filters.filter(field_type=FieldType.INT)
        chart_int_filters = SavedFilter.objects.filter(
            userID=get_user_id(request),
            parent=ModelChoices.CHART,
            chart_id=base_chart.id,
            field_type=FieldType.INT
        )
        merged_int_filters = tab_int_filters | chart_int_filters
        int_string, int_unique_filter, int_query_conditions, int_count_filters = common_integer_filter(
            merged_int_filters, return_count_filters=True
        )
        filter_string.update(int_string)

        # FLOAT FILTERS
        tab_float_filters = tab.saved_filters.filter(field_type=FieldType.FLOAT)
        chart_float_filters = SavedFilter.objects.filter(
            userID=get_user_id(request),
            parent=ModelChoices.CHART,
            chart_id=base_chart.id,
            field_type=FieldType.FLOAT
        )
        merged_float_filters = tab_float_filters | chart_float_filters
        float_string, float_unique_filter, float_query_conditions, float_count_filters = common_float_filter(
            merged_float_filters, return_count_filters=True
        )
        filter_string.update(float_string)


        base_queryset = chart_model.objects.filter(combined_q_objects).filter(**filter_string).filter(**snapshot_filter).filter(**pre_filters)
        order_by = tab.order_by
        if 'similarity' in order_by:
            order_by = 'pk'

        queryset = base_queryset
        if hasattr(chart_model, 'parent'):
            queryset = queryset.filter(parent=None)
            try:
                queryset = queryset.filter(action_status=ActionStatus.IS_ACTIVE)
            except:
                pass
        
        if user_query_conditions:
            queryset = queryset.filter(user_query_conditions)
        if date_query_conditions:
            queryset = queryset.filter(date_query_conditions)
        if int_query_conditions:
            queryset = queryset.filter(int_query_conditions)
        if float_query_conditions:
            queryset = queryset.filter(float_query_conditions)

        if user_count_filters:
            queryset = common_count_filter(user_count_filters, base_queryset, queryset, ordered_fields)
        elif date_count_filters:
            queryset = common_count_filter(date_count_filters, base_queryset, queryset, ordered_fields)
        elif int_count_filters:
            queryset = common_count_filter(int_count_filters, base_queryset, queryset, ordered_fields)
        elif float_count_filters:
            queryset = common_count_filter(float_count_filters, base_queryset, queryset, ordered_fields)
        else:
            if order_by in ['count', '-count']:
                order_by = 'pk'

        order_by = order_by or 'pk'
        queryset = queryset.order_by(order_by)

        table_name = f"{tab.content_type.app_label}_{tab.content_type.model}"
        if user_unique_filter:
            queryset = common_unique_filter(request, user_unique_filter, queryset, snapshot_filter, table_name)
        if date_unique_filter:
            queryset = common_unique_filter(request, date_unique_filter, queryset, snapshot_filter, table_name)
        if int_unique_filter:
            queryset = common_unique_filter(request, int_unique_filter, queryset, snapshot_filter, table_name)
        if float_unique_filter:
            queryset = common_unique_filter(request, float_unique_filter, queryset, snapshot_filter, table_name)

        if selected_rows:
            queryset = queryset.annotate(
                order_priority=Case(
                    *[When(pk=row_id, then=0) for row_id in selected_rows],
                    default=1,
                    output_field=IntegerField(),
                )
            ).order_by('order_priority')

        chart_list = software_filter(request, queryset, ordered_fields, search_value=tab.search, search_mode=tab.search_mode)
        mode = tab.search_mode
        if mode and mode == "graic" and chart_list:
            if 'similarity' not in ordered_fields:
                ordered_fields.insert(0, 'similarity')

        all_dates = []

        for item in chart_list:
            x_value = getattr(item, chart.x_field)

            if chart.x_field in chart_model.date_fields_to_convert and x_value:
                x_value = datetime.fromtimestamp(int(x_value)).strftime("%Y-%m-%d")

            all_dates.append(x_value)

        series_data = []

        if chart.y_field:
            for item in chart_list:
                x_value = getattr(item, chart.x_field)
                if chart.x_field in chart_model.date_fields_to_convert and x_value:
                    x_value = datetime.fromtimestamp(int(x_value)).strftime("%Y-%m-%d")
                y_value = getattr(item, chart.y_field, 0)
                series_data.append({"x": x_value, "y": y_value})
        else:
            date_counts = Counter(all_dates)
            for x_value, count in date_counts.items():
                series_data.append({"x": x_value, "y": count})

        fields = get_model_fields(f'{chart_model.__name__}', tab.pre_columns)
        saved_filters = list(SavedFilter.objects.filter(
            parent=ModelChoices.TAB, tab_id=tab.id
        ).values())
        for filter in saved_filters:
            if 'created_at' in filter:
                filter['created_at'] = filter['created_at'].isoformat()

        saved_filters_json = json.dumps(saved_filters)
        x_fields = [
        f for f in db_field_names 
            if f not in chart_model.integer_fields + chart_model.float_fields + (['count'] if 'count' in db_field_names else []) + ['ID', 'loader_instance', 'json_data', 'hash_data', 'fts']
        ]
        chart_data[str(chart.id)] = {
            'name': chart.name,
            'tab_name': tab.name,
            'type': chart.chart_type.lower(),
            'x_field': chart.x_field,
            'y_field': chart.y_field,
            'x_fields': x_fields,
            'color': chart.color,
            'data': series_data,
            'db_field_names': db_field_names,
            'fields': fields,
            'tab_id': tab.id,
            'saved_filters': saved_filters_json
        }
    
    base_fields = []
    if chart_model:
        base_fields = get_model_fields(chart_model.__name__, pre_column)

    context = {
        'base_chart': base_chart,
        'tabs': Tab.objects.filter(base_view=base_chart.base_view).order_by('created_at'),
        'charts': BaseCharts.objects.filter(base_view=base_chart.base_view),
        'chart_data': json.dumps(chart_data, cls=DjangoJSONEncoder),
        'base_fields': base_fields,
        'join_model_instance': None,
        'base_saved_filters_json': base_saved_filters_json,
        'chart_ids': chart_ids
    }
    return render(request, 'apps/charts/chart_details.html', context)


def edit_chart(request, chart_id):
    chart = get_object_or_404(TabCharts, pk=chart_id)
    if request.method == 'POST':
        chart.name = request.POST.get('name', chart.name)
        chart.chart_type = request.POST.get('chart_type', chart.chart_type)
        chart.x_field = request.POST.get('x_field', chart.x_field)
        chart.y_field = request.POST.get('y_field', chart.y_field)
        chart.color = request.POST.get('color', chart.color)
        chart.save()

        return redirect(request.META.get('HTTP_REFERER'))
    
    return redirect(request.META.get('HTTP_REFERER'))


def delete_chart(request, chart_id):
    chart = get_object_or_404(TabCharts, pk=chart_id)
    chart.delete()

    return redirect(request.META.get('HTTP_REFERER'))

# Export

import base64
import textwrap
from django.http import HttpResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Inches
from io import BytesIO
from reportlab.lib.utils import ImageReader
from apps.graic.utils import get_agent, api_engine, pre_prompt_func


def format_filter(f):
    if f.field_type != "TEXT":
        if f.value_start and f.value_end:
            value_text = f"From: {f.value_start} → To: {f.value_end}"
        elif f.value_start:
            value_text = f"From: {f.value_start}"
        elif f.value_end:
            value_text = f"To: {f.value_end}"
        else:
            value_text = ""
    else:
        value_text = f.value_start or ""

    checkboxes = []
    for attr, label in [
        ("is_not", "Not"),
        ("is_null", "Null"),
        ("is_not_null", "Not Null"),
        ("is_unique", "Unique"),
        ("is_count", "Unique Count"),
    ]:
        if getattr(f, attr, False):
            checkboxes.append(label)

    if checkboxes:
        return f"{f.field_name}: {value_text} ({', '.join(checkboxes)})"
    return f"{f.field_name}: {value_text}"


def filter_text(saved_filters):
    """Return list of formatted filter strings."""
    return [format_filter(f) for f in saved_filters]


def export_chart_docx(request):
    agent_data = get_agent("chart_agent")
    current_dt = "".join(pre_prompt_func(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"))
    custom_prompt = request.POST.get('custom_prompt', "").strip()

    if request.method != "POST":
        return HttpResponse("Invalid request")

    image_data = request.POST.get("chart_image", "")
    chart_id = request.POST.get("chart_id", "")
    chart = get_object_or_404(TabCharts, id=chart_id)

    saved_filters = SavedFilter.objects.filter(
        userID=get_user_id(request),
        parent=ModelChoices.TAB,
        tab_id=chart.parent_tab.id
    )

    filter_texts = filter_text(saved_filters)
    filters_str = "\n".join(filter_texts) if filter_texts else "No filters applied"

    if not image_data or not image_data.startswith("data:image"):
        return HttpResponse("Invalid image", status=400)

    system_prompt = agent_data['system_prompt']
    if custom_prompt:
        system_prompt = f"{system_prompt}\n\n{custom_prompt}"

    full_prompt = (
        f"{current_dt}\n\n{system_prompt}\n\n"
        f"{agent_data.get('instruction', '')}\n\n"
        f"Applied Filters:\n{filters_str}"
    )

    ai_analysis = api_engine(
        prompt=full_prompt,
        is_online=True,
        llm=agent_data.get("llm", "openai"),
        image_data=image_data
    )

    _, encoded = image_data.split(",", 1)
    image_bytes = base64.b64decode(encoded)

    doc = Document()
    doc.add_heading(f"GRC {chart.name} Report - {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}", level=1)
    doc.add_picture(BytesIO(image_bytes), width=Inches(6))

    doc.add_heading("Filters", level=2)
    if filter_texts:
        for text in filter_texts:
            doc.add_paragraph(text)
    else:
        doc.add_paragraph("No filters applied.")

    doc.add_heading("grAIc Analysis", level=2)
    doc.add_paragraph(ai_analysis)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{chart.name}_{timestamp}.docx"

    export_option = request.POST.get("export_option", "download")

    if export_option == "grc":
        saved_file = get_chart_save_path(request.user, output, filename)
        return JsonResponse({
            "status": "saved",
            **saved_file
        })

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def export_chart_pdf(request):
    agent_data = get_agent("chart_agent")
    current_dt = "".join(pre_prompt_func(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"))
    custom_prompt = request.POST.get('custom_prompt', "").strip()

    if request.method != "POST":
        return HttpResponse("Invalid request")

    image_data = request.POST.get("chart_image", "")
    chart_id = request.POST.get("chart_id", "")
    chart = get_object_or_404(TabCharts, id=chart_id)

    if not image_data or not image_data.startswith("data:image"):
        return HttpResponse("Invalid image", status=400)

    saved_filters = SavedFilter.objects.filter(
        userID=get_user_id(request),
        parent=ModelChoices.TAB,
        tab_id=chart.parent_tab.id
    )

    filter_texts = filter_text(saved_filters)
    filters_str = "\n".join(filter_texts) if filter_texts else "No filters applied"

    system_prompt = agent_data['system_prompt']
    if custom_prompt:
        system_prompt = f"{system_prompt}\n\n{custom_prompt}"

    full_prompt = (
        f"{current_dt}\n\n{system_prompt}\n\n"
        f"{agent_data.get('instruction', '')}\n\n"
        f"Applied Filters:\n{filters_str}"
    )

    ai_analysis = api_engine(
        prompt=full_prompt,
        is_online=True,
        llm=agent_data.get("llm", "openai"),
        image_data=image_data
    )

    _, encoded = image_data.split(",", 1)
    image_bytes = base64.b64decode(encoded)

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    img = ImageReader(BytesIO(image_bytes))
    p.drawImage(img, 50, height - 400, width=500, preserveAspectRatio=True, mask='auto')

    p.setFont("Helvetica-Bold", 18)
    p.drawString(50, height - 40, f"GRC {chart.name} Report - {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}")

    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, height - 420, "Filters:")
    p.setFont("Helvetica", 12)
    y_pos = height - 440
    for text in filter_texts or ["No filters applied."]:
        wrapped_lines = textwrap.wrap(text, width=90)
        for line in wrapped_lines:
            p.drawString(50, y_pos, line)
            y_pos -= 15

    y_pos -= 10
    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, y_pos, "grAIc Analysis:")
    y_pos -= 20
    p.setFont("Helvetica", 12)
    for line in ai_analysis.split("\n"):
        for wrapped_line in textwrap.wrap(line, width=90):
            p.drawString(50, y_pos, wrapped_line)
            y_pos -= 15
            if y_pos < 50:
                p.showPage()
                y_pos = height - 50
                p.setFont("Helvetica", 12)

    p.showPage()
    p.save()

    pdf = buffer.getvalue()
    
    timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{chart.name}_{timestamp}.pdf"

    export_option = request.POST.get("export_option", "download")

    if export_option == "grc":
        saved_file = get_chart_save_path(request.user, BytesIO(pdf), filename)
        return JsonResponse({
            "status": "saved",
            **saved_file
        })

    buffer.close()
    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def export_charts_docx_bulk(request):
    if request.method != "POST":
        return HttpResponse("Invalid request", status=400)

    chart_ids = request.POST.get("chart_ids", "[]")
    try:
        chart_ids = json.loads(chart_ids)
    except:
        chart_ids = []

    agent_data = get_agent("chart_agent")
    current_dt = "".join(pre_prompt_func(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"))

    doc = Document()
    doc.add_heading(f"GRC Chart Report - {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}", level=1)
    custom_prompt = request.POST.get('custom_prompt', "").strip()

    for chart_id in chart_ids:
        chart = get_object_or_404(TabCharts, id=chart_id)
        saved_filters = SavedFilter.objects.filter(
            userID=get_user_id(request),
            parent=ModelChoices.TAB,
            tab_id=chart.parent_tab.id
        )

        img_data = request.POST.get(f"chart_image_{chart_id}", "")
        if img_data.startswith("data:image"):
            _, encoded = img_data.split(",", 1)
            image_bytes = base64.b64decode(encoded)
            img_buffer = BytesIO(image_bytes)

            doc.add_heading(f"{chart.name}", level=2)
            doc.add_picture(img_buffer, width=Inches(6))

        filter_texts = filter_text(saved_filters)
        filters_str = "\n".join(filter_texts) if filter_texts else "No filters applied"

        doc.add_heading("Filters", level=2)
        if filter_texts:
            for text in filter_texts:
                doc.add_paragraph(text)
        else:
            doc.add_paragraph("No filters applied.")

        system_prompt = agent_data['system_prompt']
        if custom_prompt:
            system_prompt = f"{system_prompt}\n\n{custom_prompt}"
        
        full_prompt = (
            f"{current_dt}\n\n{system_prompt}\n\n"
            f"{agent_data.get('instruction', '')}\n\n"
            f"Applied Filters:\n{filters_str}"
        )

        ai_analysis = api_engine(
            prompt=full_prompt,
            is_online=True,
            llm=agent_data.get("llm", "openai"),
            image_data=img_data
        )

        doc.add_heading("grAIc Analysis", level=2)
        doc.add_paragraph(ai_analysis)

        doc.add_page_break()

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")
    filename = f"GRC_Chart_Report_{timestamp}.docx"

    export_option = request.POST.get("export_option", "download")

    if export_option == "grc":
        saved_file = get_chart_save_path(request.user, output, filename)
        return JsonResponse({
            "status": "saved",
            **saved_file
        })

    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def export_charts_pdf_bulk(request):
    if request.method != "POST":
        return HttpResponse("Invalid request", status=400)

    chart_ids = request.POST.get("chart_ids", "[]")
    try:
        chart_ids = json.loads(chart_ids)
    except:
        chart_ids = []

    agent_data = get_agent("chart_agent")
    current_dt = "".join(pre_prompt_func(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"))

    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    p.setFont("Helvetica-Bold", 22)
    p.drawString(
        50,
        height - 40,
        f"GRC Chart Report - {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    custom_prompt = request.POST.get('custom_prompt', "").strip()

    for chart_id in chart_ids:
        chart = get_object_or_404(TabCharts, id=chart_id)
        saved_filters = SavedFilter.objects.filter(
            userID=get_user_id(request),
            parent=ModelChoices.TAB,
            tab_id=chart.parent_tab.id
        )

        y_position = height - 80
        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, y_position, chart.name)

        img_data = request.POST.get(f"chart_image_{chart_id}", "")
        if img_data.startswith("data:image"):
            _, encoded = img_data.split(",", 1)
            image_bytes = base64.b64decode(encoded)
            img = ImageReader(BytesIO(image_bytes))
            p.drawImage(img, 50, y_position - 350, width=500, preserveAspectRatio=True, mask='auto')

        y_position -= 380

        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, y_position, "Filters:")
        y_position -= 20
        p.setFont("Helvetica", 12)

        filter_texts = filter_text(saved_filters)
        filters_str = "\n".join(filter_texts) if filter_texts else "No filters applied"

        if filter_texts:
            for text in filter_texts:
                for line in textwrap.wrap(text, width=90):
                    p.drawString(60, y_position, line)
                    y_position -= 15
                    if y_position < 50:
                        p.showPage()
                        y_position = height - 50
                        p.setFont("Helvetica", 12)
        else:
            p.drawString(60, y_position, "No filters applied.")
            y_position -= 20

        system_prompt = agent_data['system_prompt']
        if custom_prompt:
            system_prompt = f"{system_prompt}\n\n{custom_prompt}"

        full_prompt = (
            f"{current_dt}\n\n{agent_data['system_prompt']}\n\n"
            f"{agent_data.get('instruction', '')}\n\n"
            f"Applied Filters:\n{filters_str}"
        )

        ai_analysis = api_engine(
            prompt=full_prompt,
            is_online=True,
            llm=agent_data.get("llm", "openai"),
            image_data=img_data
        )

        y_position -= 10
        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, y_position, "grAIc Analysis:")
        y_position -= 20
        p.setFont("Helvetica", 12)
        for line in ai_analysis.split("\n"):
            for wrapped_line in textwrap.wrap(line, width=90):
                p.drawString(50, y_position, wrapped_line)
                y_position -= 15
                if y_position < 50:
                    p.showPage()
                    y_position = height - 50
                    p.setFont("Helvetica", 12)

        p.showPage()

    p.save()
    pdf = buffer.getvalue()

    timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")
    filename = f"GRC_Chart_Report_{timestamp}.pdf"

    export_option = request.POST.get("export_option", "download")

    if export_option == "grc":
        saved_file = get_chart_save_path(request.user, BytesIO(pdf), filename)
        return JsonResponse({
            "status": "saved",
            **saved_file
        })

    buffer.close()
    response = HttpResponse(pdf, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response

def get_or_create_folder(name, parent=None, user=None):
    folder, _ = FileManager.objects.get_or_create(
        name=name,
        parent=parent,
        defaults={
            "created_by": user,
            "updated_by": user,
            "action_status": ActionStatus.IS_ACTIVE,
            "folder_status": DocumentStatus.APPROVED,
        }
    )
    return folder


def get_chart_save_path(user, buffer, filename):
    buffer.seek(0)
    username = user.username if user.is_authenticated else "anonymous"

    users_root = get_or_create_folder("users", parent=None, user=user)
    user_folder = get_or_create_folder(username, parent=users_root, user=user)
    chart_folder = get_or_create_folder("chart", parent=user_folder, user=user)

    file_obj = File.objects.create(
        file_manager=chart_folder,
        file=ContentFile(buffer.getvalue(), name=filename),
        action_status=ActionStatus.IS_ACTIVE,
        file_status=DocumentStatus.APPROVED,
        uploaded_by=user,
        updated_by=user
    )

    base_path = os.path.join(settings.MEDIA_ROOT, "users", username, "chart")
    os.makedirs(base_path, exist_ok=True)
    file_path = os.path.join(base_path, filename)
    with open(file_path, "wb") as f:
        f.write(buffer.getvalue())

    return {
        "file_id": file_obj.id,
        "filename": filename,
        "path": f"users/{username}/chart"
    }



#

from django.views.decorators.http import require_POST
from django.utils.dateparse import parse_datetime
from apps.tables.models import ScheduledChartExport

@require_POST
def create_chart_export_schedule(request):
    chart = get_object_or_404(TabCharts, id=request.POST.get("chart_id"))

    weekdays = request.POST.getlist("weekdays[]")
    chart_image = request.POST.get("chart_image", "")

    schedule = ScheduledChartExport.objects.create(
        user=request.user,
        chart=chart,
        export_type=request.POST.get("export_type"),
        frequency=request.POST.get("frequency"),
        chart_image=chart_image,
        start_at=parse_datetime(request.POST.get("start_at")),
        end_at=parse_datetime(request.POST.get("end_at")) if request.POST.get("end_at") else None,
        hour_interval=request.POST.get("hour_interval") or None,
        time_of_day=request.POST.get("time_of_day") or None,
        month_day=request.POST.get("month_day") or None,
        weekdays=weekdays,
        custom_prompt=request.POST.get("custom_prompt", ""),
        export_option="grc"
    )

    return JsonResponse({
        "status": "scheduled",
        "schedule_id": schedule.id
    })


from celery import shared_task

def generate_chart_image(schedule: ScheduledChartExport):
    if not schedule.chart_image:
        raise ValueError("No chart image stored for this scheduled export")
    return schedule.chart_image


def generate_chart_export(
    *,
    chart_id,
    schedule_id,
    user,
    export_type,
    custom_prompt="",
    export_option="grc",
):
    chart = get_object_or_404(TabCharts, id=chart_id)
    schedule = get_object_or_404(ScheduledChartExport, id=schedule_id)

    image_data = generate_chart_image(schedule)
    if not image_data.startswith("data:image"):
        raise ValueError("Invalid generated image")

    agent_data = get_agent("chart_agent")
    current_dt = "".join(
        pre_prompt_func(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    )

    saved_filters = SavedFilter.objects.filter(
        userID=user.id,
        parent=ModelChoices.TAB,
        tab_id=chart.parent_tab.id,
    )

    filter_texts = filter_text(saved_filters)
    filters_str = "\n".join(filter_texts) if filter_texts else "No filters applied"

    system_prompt = agent_data["system_prompt"]
    if custom_prompt:
        system_prompt = f"{system_prompt}\n\n{custom_prompt}"

    full_prompt = (
        f"{current_dt}\n\n{system_prompt}\n\n"
        f"{agent_data.get('instruction', '')}\n\n"
        f"Applied Filters:\n{filters_str}"
    )

    ai_analysis = api_engine(
        prompt=full_prompt,
        is_online=True,
        llm=agent_data.get("llm", "openai"),
        image_data=image_data,
    )

    _, encoded = image_data.split(",", 1)
    image_bytes = base64.b64decode(encoded)

    timestamp = timezone.now().strftime("%Y%m%d_%H%M%S")

    if export_type == "docx":
        return _build_docx(
            chart, image_bytes, filter_texts, ai_analysis, user, export_option, timestamp
        )

    return _build_pdf(
        chart, image_bytes, filter_texts, ai_analysis, user, export_option, timestamp
    )


def _build_docx(chart, image_bytes, filter_texts, ai_analysis, user, export_option, timestamp):
    doc = Document()
    doc.add_heading(f"GRC {chart.name} Report - {timestamp}", level=1)
    doc.add_picture(BytesIO(image_bytes), width=Inches(6))

    doc.add_heading("Filters", level=2)
    for text in filter_texts or ["No filters applied."]:
        doc.add_paragraph(text)

    doc.add_heading("grAIc Analysis", level=2)
    doc.add_paragraph(ai_analysis)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"{chart.name}_{timestamp}.docx"

    if export_option == "grc":
        return get_chart_save_path(user, output, filename)

    return output.getvalue(), filename

def _build_pdf(chart, image_bytes, filter_texts, ai_analysis, user, export_option, timestamp):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    img = ImageReader(BytesIO(image_bytes))
    p.drawImage(img, 50, height - 400, width=500, preserveAspectRatio=True)

    p.setFont("Helvetica-Bold", 18)
    p.drawString(50, height - 40, f"GRC {chart.name} Report - {timestamp}")

    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, height - 420, "Filters:")

    p.setFont("Helvetica", 12)
    y = height - 440
    for text in filter_texts or ["No filters applied."]:
        for line in textwrap.wrap(text, 90):
            p.drawString(50, y, line)
            y -= 15

    y -= 10
    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, y, "grAIc Analysis:")
    y -= 20

    p.setFont("Helvetica", 12)
    for line in ai_analysis.split("\n"):
        for wrapped in textwrap.wrap(line, 90):
            if y < 50:
                p.showPage()
                y = height - 50
            p.drawString(50, y, wrapped)
            y -= 15

    p.save()
    pdf = buffer.getvalue()

    filename = f"{chart.name}_{timestamp}.pdf"

    if export_option == "grc":
        return get_chart_save_path(user, BytesIO(pdf), filename)

    return pdf, filename


@shared_task
def run_scheduled_chart_exports():
    now = timezone.now()

    schedules = ScheduledChartExport.objects.filter(
        is_active=True,
        start_at__lte=now,
    )

    for schedule in schedules:
        if schedule.end_at and now > schedule.end_at:
            schedule.is_active = False
            schedule.save(update_fields=["is_active"])
            continue

        if not should_run_now(schedule, now):
            continue

        generate_chart_export(
            chart_id=schedule.chart_id,
            schedule_id=schedule.id,
            user=schedule.user,
            export_type=schedule.export_type,
            custom_prompt=schedule.custom_prompt,
            export_option="grc",
        )

        schedule.last_run_at = now
        if schedule.frequency == "once":
            schedule.is_active = False

        schedule.save(update_fields=["last_run_at", "is_active"])


def should_run_now(schedule, now):
    if schedule.frequency == "once":
        return schedule.last_run_at is None

    if schedule.frequency == "hourly":
        if not schedule.last_run_at:
            return True
        return (now - schedule.last_run_at).total_seconds() >= schedule.hour_interval * 3600

    if schedule.frequency == "daily":
        return (
            now.hour == schedule.time_of_day.hour and
            schedule.last_run_at is None or schedule.last_run_at.date() < now.date()
        )

    if schedule.frequency == "weekly":
        return (
            now.weekday() in schedule.weekdays and
            now.hour == schedule.time_of_day.hour and
            (not schedule.last_run_at or schedule.last_run_at.date() < now.date())
        )

    if schedule.frequency == "monthly":
        return (
            now.day == schedule.month_day and
            now.hour == schedule.time_of_day.hour and
            (not schedule.last_run_at or schedule.last_run_at.month < now.month)
        )

    return False
