import re
import numpy as np
from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse
from apps.graic.models import PrePrompt, Graic, Agent, Chat
from datetime import datetime
from apps.common.signals import to_sql_vector
from apps.file_manager.models import File, ActionStatus, FileChunk, DefaultValues, FilterMethod, FileManager
from loader.models import is_pgvector_enabled, InstantUpload
from apps.graic.utils import pre_prompt_func, api_engine, aggregate_responses, extract_keywords, extract_keywords_llm, get_last_responses_summary, get_agent, pick_agent_with_llm
from django.apps import apps
from django.db import connection
from django.contrib.contenttypes.models import ContentType
from django.urls import reverse
from loader.models import preprocess_text
from openpyxl import load_workbook
from tempfile import NamedTemporaryFile
from openpyxl.utils import get_column_letter
from django.db.models import Q
from apps.tables.models import ActionStatus, DocumentStatus, DocumentType

GRAIC_SEQUENCE_PATTERN = re.compile(
    r'<graic\s+sequence=["\']?(\d+)["\']?(?:\s+new_chat=["\']?(\w+)["\']?)?>\s*(.*?)\s*</graic>',
    re.DOTALL | re.IGNORECASE
)

# Create your views here.

def graic_chat_index(request):
    default_value = DefaultValues.objects.first() or DefaultValues.objects.create()

    if request.method == 'POST':
        prompt = request.POST.get('prompt')
        if not prompt:
            return render(request, 'graic/chat_index.html', {"error": "Prompt cannot be empty"})

        chat = Chat.objects.create(
            user=request.user,
            title=prompt[:150]
        )

        start_time = datetime.now()
        prompt = request.POST.get('prompt')
        uploaded_file = request.FILES.get('file')
        existing_file_id = request.POST.get('existing_file_id')

        agent = pick_agent_with_llm(prompt)
        print(agent, "===agent===")

        generated_file_path = None
        summary_response = ""
        pre_response = ""

        print(uploaded_file, "===file===")

        if uploaded_file:     
            with NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_in, \
                NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_out:

                for chunk in uploaded_file.chunks():
                    tmp_in.write(chunk)

                process_excel_with_sequences(request, tmp_in.name, tmp_out.name, chat.uuid, default_value)
                generated_file_path = tmp_out.name

        elif existing_file_id:
            file_instance = get_object_or_404(File, id=existing_file_id)
            with NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_in, \
                NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_out:

                with file_instance.file.open('rb') as f:
                    tmp_in.write(f.read())

                process_excel_with_sequences(request, tmp_in.name, tmp_out.name, chat.uuid, default_value)
                generated_file_path = tmp_out.name

        elif agent == 'file_agent':
            pre_response = file_process(request, default_value, chat.uuid, agent)
        elif agent == 'sql_agent':
            pre_response = sql_process(request, default_value, chat.uuid, agent)
        else:
            pre_response = dt_process(request, default_value, chat.uuid, agent)

        is_dt = False
        summary_response = ""
        end_time = datetime.now()
        html_response = aggregate_responses(
            final_response=pre_response,
            prompt=prompt,
            user_id=request.user.id,
            chat_id=chat.uuid,
            metadata={
                'time_taken': (end_time - start_time).total_seconds(),
                'is_dt': is_dt,
                'summary_response': summary_response
            },
            generated_file_path=generated_file_path
        )

        return JsonResponse({
            "status": "complete",
            "response": html_response,
            "is_dt": is_dt,
            "chat_id": chat.uuid,
            'time_taken': (end_time - start_time).total_seconds()
        })

        # return redirect(reverse('graic_chat', args=[chat.uuid]))

    return render(request, 'graic/chat_index.html')

def graic_chat(request, chat_id):
    chat = get_object_or_404(Chat, uuid=chat_id)
    griac_data = Graic.objects.filter(user=request.user, chat=chat)

    context = {
        'griac_data': griac_data,
        'chat': chat,
        'segment': chat.pk,
    }
    return render(request, 'graic/chat.html', context)


def file_process(request, default_value, chat_id, agent="summary_agent", prompt="", previous_response="", previous_prompt="", new_chat="N"):
    if not prompt:
        prompt = request.POST.get('prompt')

    client_datetime = request.POST.get('client_datetime')

    files = list(File.objects.filter(
        action_status=ActionStatus.IS_ACTIVE
    ).order_by('-uploaded_at'))

    chunk_texts = []

    if prompt and is_pgvector_enabled():
        if default_value and default_value.extract_keywords:
            keywords = extract_keywords_llm(prompt)
        else:
            keywords = prompt

        print(keywords, "====keywords====")

        prompt_vector = to_sql_vector(keywords)

        all_chunk_similarities = []
        file_chunks_map = {}

        for file in files:
            chunks = FileChunk.objects.filter(file=file).exclude(vector=None)
            if not chunks.exists():
                continue

            chunk_similarities = []
            for chunk in chunks:
                similarity = np.dot(np.array(chunk.vector), np.array(prompt_vector)) / (
                    np.linalg.norm(chunk.vector) * np.linalg.norm(prompt_vector)
                )
                similarity = round(float(similarity), 3)
                chunk_similarities.append((similarity, chunk))

            file_chunks_map[file] = chunk_similarities
            all_chunk_similarities.extend(chunk_similarities)

        if default_value and default_value.dynamic_similarity and default_value.dynamic_rows:
            sorted_all_chunks = sorted(all_chunk_similarities, key=lambda x: x[0], reverse=True)
            if len(sorted_all_chunks) >= int(default_value.dynamic_rows):
                nth_index = int(default_value.dynamic_rows) - 1
                nth_similarity = sorted_all_chunks[nth_index][0]
            else:
                nth_similarity = sorted_all_chunks[0][0]
        else:
            nth_similarity = None

        matched_chunks = []
        for file, chunk_similarities in file_chunks_map.items():
            for sim, chunk in chunk_similarities:
                if default_value and default_value.dynamic_similarity and nth_similarity is not None:
                    if sim >= nth_similarity:
                        matched_chunks.append((sim, chunk))
                else:
                    similarity_threshold = (
                        int(default_value.similarity_threshold) / 100
                        if default_value and default_value.similarity_threshold
                        else 0.2
                    )
                    if sim >= similarity_threshold:
                        matched_chunks.append((sim, chunk))

        matched_chunks = sorted(matched_chunks, key=lambda x: x[0], reverse=True)
        top_chunks = matched_chunks[:10]

        for sim, chunk in matched_chunks:
            text = chunk.chunk_text or ""
            chunk_texts.append({
                "file_id": chunk.file_id,
                'file_name': chunk.file.file_name(),
                'file_path': request.build_absolute_uri(chunk.file.file.url) if chunk.file.file else '',
                "page": chunk.page_number,
                "similarity": sim,
                "text": text.strip()
            })

    agent_data  = get_agent(agent)
    conversation_summary = get_last_responses_summary(request.user.id, chat_id, is_online=agent_data['online'], llm=agent_data['llm'])

    if agent in "general_agent":
        ai_prompt = (
            "".join(pre_prompt_func(client_datetime)) + "\n\n"
            f"{agent_data['system_prompt']}\n\n"
            f"{agent_data['instruction']}\n\n"
            f"Conversation Summary:\n{conversation_summary}\n\n"
            f"User Prompt: {prompt}\n\n"
        )
    else:
        ai_prompt = (
            "".join(pre_prompt_func(client_datetime)) + "\n\n"
            f"{agent_data['system_prompt']}\n\n"
            f"{agent_data['instruction']}\n\n"
            f"Conversation Summary:\n{conversation_summary}\n\n"
            f"User Prompt: {prompt}\n\n"
            "Context Chunks:\n"
            + "\n".join(
                f"[File ID: {chunk['file_id']}, File Name: {chunk['file_name']}, File Path: {chunk['file_path']}, Page: {chunk['page']}, Similarity: {chunk['similarity']}]\n{chunk['text']}"
                for chunk in chunk_texts
            )
        )

    print("\n============== AI Prompt start ================\n")
    print(ai_prompt)
    print("\n============== AI Prompt end ================\n")

    pre_response = api_engine(ai_prompt, agent_data['online'], agent_data['llm'])

    return pre_response

def estimate_tokens(text):
    return len(text) / 4
    
def rotate_context_window(ai_prompt_base, chunk_texts, agent_data, max_batch_tokens=12000):
    summary = ""
    batch = []
    current_tokens = estimate_tokens(ai_prompt_base)

    for chunk in chunk_texts:
        chunk_text = f"[Table: {chunk['table']}, ID: {chunk['id']}, Similarity: {chunk['similarity']}]\n{preprocess_text(chunk['json'])}\n\n"
        tokens = estimate_tokens(chunk_text)

        if estimate_tokens(summary) > max_batch_tokens * 0.4:
            compression_prompt = f"Compress the following summary while keeping all key factual details:\n\n{summary}"
            summary = api_engine(compression_prompt, agent_data['online'], agent_data['llm'])

        if current_tokens + tokens > max_batch_tokens:
            ai_prompt = (
                ai_prompt_base +
                f"\n\nPrevious Summary:\n{summary}\n\n" +
                "Context Batch:\n" + "".join(batch)
            )
            response = api_engine(ai_prompt, agent_data['online'], agent_data['llm'])
            summary = response
            batch = []
            current_tokens = estimate_tokens(ai_prompt_base)
        batch.append(chunk_text)
        current_tokens += tokens

    if batch:
        ai_prompt = (
            ai_prompt_base +
            f"\n\nPrevious Summary:\n{summary}\n\n" +
            "Context Batch:\n" + "".join(batch)
        )
        response = api_engine(ai_prompt, agent_data['online'], agent_data['llm'])
        summary = response

    return summary


def dt_process(request, default_value, chat_id, agent="summary_agent", prompt="", previous_response="", previous_prompt="", new_chat="N"):
    if not prompt:
        prompt = request.POST.get('prompt')

    client_datetime = request.POST.get('client_datetime')

    chunk_texts = []

    if prompt and is_pgvector_enabled():
        if default_value and default_value.extract_keywords:
            keywords = extract_keywords_llm(prompt)
        else:
            keywords = prompt

        keywords = preprocess_text(keywords)
        print(keywords, "====keywords====")
        prompt_vector = to_sql_vector(keywords)
        tables_app_models = apps.get_app_config('tables').get_models()
        models_with_fts = [
            model for model in tables_app_models
            if 'fts' in [f.name for f in model._meta.fields]
        ]

        matched_chunks = []

        for model in models_with_fts:
            if hasattr(model, 'hash_data'):
                table_name = model._meta.db_table
                content_type = ContentType.objects.get_for_model(model)
                snapshots = InstantUpload.objects.filter(content_type=content_type).order_by('-created_at')
                snapshot = request.GET.get('snapshot')
                snapshot_value = None
                if snapshot and snapshot != 'all':
                    try:
                        summary = InstantUpload.objects.get(id=snapshot, content_type=content_type)
                        snapshot_value = summary.pk
                    except InstantUpload.DoesNotExist:
                        snapshot_value = None

                elif snapshots.exists():
                    latest_snapshot = snapshots.first()
                    snapshot_value = latest_snapshot.pk

                excluded_fields = {"loader_instance", "hash_data", "fts"}
                fields = [f.name for f in model._meta.fields if f.name not in excluded_fields]

                field_list = ', '.join([f't."{field}"' for field in fields])

                query = f"""
                    WITH q AS (
                        SELECT '{prompt_vector}'::vector AS v
                    )
                    SELECT 
                        t."ID",
                        {field_list},
                        ROUND(1 - (t.hash_data <=> q.v)::numeric, 3) AS similarity
                    FROM public."{table_name}" t, q
                    WHERE t.loader_instance IS NOT NULL
                """
                if snapshot_value:
                    query += f" AND t.loader_instance = {snapshot_value}"

                query += " ORDER BY t.hash_data <=> q.v;"

                with connection.cursor() as cursor:
                    cursor.execute(query)
                    results = cursor.fetchall()

                for row in results:
                    row_dict = dict(zip(["id"] + fields + ["similarity"], row))

                    matched_chunks.append({
                        "table": table_name,
                        "id": row_dict["id"],
                        "similarity": row_dict["similarity"],
                        "json": row_dict["json_data"],
                    })

        matched_chunks = sorted(matched_chunks, key=lambda x: x["similarity"], reverse=True)

        if default_value and default_value.dynamic_similarity and default_value.dynamic_rows:
            if matched_chunks:
                rows_n = int(default_value.dynamic_rows)

                similarities_percent = [
                    c["similarity"] * 100 if c["similarity"] < 1 else c["similarity"]
                    for c in matched_chunks
                ]

                grouped = sorted({int(s) for s in similarities_percent}, reverse=True)

                if len(grouped) >= rows_n:
                    nth_group = grouped[rows_n - 1]
                else:
                    nth_group = grouped[-1]

                matched_chunks = [
                    c for c in matched_chunks
                    if int(c["similarity"] * 100 if c["similarity"] < 1 else c["similarity"]) >= nth_group
                ]
        else:
            rows_n = 10

            similarities_percent = [
                c["similarity"] * 100 if c["similarity"] < 1 else c["similarity"]
                for c in matched_chunks
            ]

            grouped = sorted({int(s) for s in similarities_percent}, reverse=True)

            if len(grouped) >= rows_n:
                nth_group = grouped[rows_n - 1]
            else:
                nth_group = grouped[-1]

            matched_chunks = [
                c for c in matched_chunks
                if int(c["similarity"] * 100 if c["similarity"] < 1 else c["similarity"]) >= nth_group
            ]

        matched_chunks = sorted(matched_chunks, key=lambda x: x["similarity"], reverse=True)

        chunk_texts.extend(matched_chunks)

    agent_data  = get_agent(agent)
    conversation_summary = get_last_responses_summary(request.user.id, chat_id, is_online=agent_data['online'], llm=agent_data['llm'])

    if agent in "general_agent":
        if previous_response and previous_prompt:
            ai_prompt = (
                "".join(pre_prompt_func(client_datetime)) + "\n\n"
                f"{agent_data['system_prompt']}\n"
                f"{agent_data['instruction']}\n"
            )
            if new_chat == 'N':
                ai_prompt += (
                    f"Previous Prompt:\n{previous_prompt}\n"
                    f"Previous Response:\n{previous_response}\n"
                )

            ai_prompt += f"User Prompt: {prompt}\n\n"

        else:
            ai_prompt = (
                "".join(pre_prompt_func(client_datetime))
                + "\n\n"
                + agent_data['system_prompt']
                + "\n\n"
                + agent_data['instruction']
                + "\n\n"
            )
            if new_chat == 'N':
                ai_prompt += f"Conversation Summary:\n{conversation_summary}\n\n"

            ai_prompt += f"User Prompt: {prompt}\n\n"

    else:
        if previous_response and previous_prompt:
            ai_prompt = (
                "".join(pre_prompt_func(client_datetime)) + "\n\n"
                f"{agent_data['system_prompt']}\n"
                f"{agent_data['instruction']}\n"
            )
            if new_chat == 'N':
                ai_prompt += (
                    f"Previous Prompt:\n{previous_prompt}\n"
                    f"Previous Response:\n{previous_response}\n"
                )

            ai_prompt += f"User Prompt: {prompt}\n\n"

        else:
            ai_prompt = "".join(pre_prompt_func(client_datetime)) + "\n\n"
            ai_prompt += f"{agent_data['system_prompt']}\n\n"
            ai_prompt += f"{agent_data['instruction']}\n\n"

            if new_chat == 'N':
                ai_prompt += f"Conversation Summary:\n{conversation_summary}\n\n"

            ai_prompt += f"User Prompt: {prompt}\n\n"
            ai_prompt += "Context Chunks:\n" + "\n".join(
                f"[Table: {chunk['table']}, ID: {chunk['id']}, Similarity: {chunk['similarity']}]\n"
                f"{preprocess_text(chunk['json'])}"
                for chunk in chunk_texts
            )

    print("\n============== AI Prompt start ================\n")
    print(ai_prompt)
    print("\n============== AI Prompt end ================\n")

    MAX_USAGE_RATIO = 0.85
    max_context_tokens = 128000

    base_tokens = estimate_tokens(ai_prompt)
    chunk_tokens = sum(estimate_tokens(preprocess_text(chunk['json'])) for chunk in chunk_texts)

    total_tokens = base_tokens + chunk_tokens

    if total_tokens > max_context_tokens * MAX_USAGE_RATIO:
        pre_response = rotate_context_window(
            ai_prompt,
            chunk_texts,
            agent_data,
            max_batch_tokens=int(max_context_tokens * MAX_USAGE_RATIO)
        )
    else:
        pre_response = api_engine(ai_prompt, agent_data['online'], agent_data['llm'])

    return pre_response


def sql_process(request, default_value, chat_id, agent="sql_agent", prompt="", previous_response="", previous_prompt="", new_chat="N"):
    if not prompt:
        prompt = request.POST.get('prompt')

    client_datetime = request.POST.get('client_datetime')

    print(prompt, "====keywords====")

    tables_app_models = apps.get_app_config('tables').get_models()
    models_with_fts = [
        model for model in tables_app_models
        if any(f.name == 'fts' for f in model._meta.fields)
    ]

    if not models_with_fts:
        return "No models with FTS found."

    agent_data = get_agent(agent)
    conversation_summary = get_last_responses_summary(
        request.user.id,
        chat_id,
        is_online=agent_data['online'],
        llm=agent_data['llm']
    )

    dataset_schema = []
    for model in models_with_fts:
        if hasattr(model, 'hash_data'):
            table_name = model._meta.db_table
            content_type = ContentType.objects.get_for_model(model)

            snapshots = InstantUpload.objects.filter(content_type=content_type).order_by('-created_at')
            snapshot_value = None
            snapshot = request.GET.get('snapshot')
            if snapshot and snapshot != 'all':
                snapshot_value = InstantUpload.objects.filter(
                    id=snapshot, content_type=content_type
                ).values_list('pk', flat=True).first()
            elif snapshots.exists():
                snapshot_value = snapshots.first().pk

            excluded_fields = {"loader_instance", "hash_data", "fts"}
            fields = [f.name for f in model._meta.fields if f.name not in excluded_fields]

            schema_text = (
                f"Table: {table_name}\n"
                f"Columns: {', '.join(fields)}\n"
                f"Snapshot ID: {snapshot_value or 'None'}\n"
            )
            dataset_schema.append(schema_text)


    ai_prompt = (
        "".join(pre_prompt_func(client_datetime)) + "\n\n"
        f"{agent_data['system_prompt']}\n\n"
        f"{agent_data['instruction']}\n\n"
        "Important:\n"
        "- You may generate multiple SELECT statements — one per table that could contain relevant data.\n"
        "- Each query must target only one table.\n"
        "- Do NOT combine tables using UNION, UNION ALL, or JOIN.\n"
        "- Separate each query with a semicolon (;).\n"
        "- Return only raw SQL queries, no explanation.\n\n"
        f"Conversation Summary:\n{conversation_summary}\n\n"
        f"User Prompt:\n{prompt}\n\n"
        "Dataset Schema:\n" + "\n".join(dataset_schema)
    )


    print("\n============== AI Prompt start ================\n")
    print(ai_prompt)
    print("\n============== AI Prompt end ================\n")

    sql_response = api_engine(ai_prompt, agent_data['online'], agent_data['llm'])
    print("Generated SQL Response:\n", sql_response)

    sql_queries = [q.strip() for q in re.split(r';\s*', sql_response) if q.strip().lower().startswith('select')]

    query_results = []

    for i, query in enumerate(sql_queries):
        print(f"Executing query {i+1}/{len(sql_queries)}")
        try:
            with connection.cursor() as cursor:
                cursor.execute(query)
                print(f"Query {i+1} executed successfully")

                columns = [col[0] for col in cursor.description]
                if "json_data" not in columns:
                    table_result = []
                else:
                    idx = columns.index("json_data")
                    table_result = [row[idx] for row in cursor.fetchall()]

                match = re.search(r'from\s+([a-zA-Z0-9_"]+)', query, re.IGNORECASE)
                table_name = match.group(1).replace('"', '') if match else "unknown_table"

                query_results.append({
                    "table": table_name,
                    "query": query,
                    "rows": table_result
                })

        except Exception as e:
            print(f"Error in query {i+1}: {e}")
            error_message = str(e)

            fix_prompt = (
                f"You are a PostgreSQL expert.\n"
                f"The following SQL query caused an error.\n"
                f"Error: {error_message}\n"
                f"Query: {query}\n\n"
                "Please fix the query. Only return the corrected SQL (no explanation)."
            )
            fixed_query = api_engine(fix_prompt, agent_data['online'], agent_data['llm']).strip()
            print(f"Retrying with fixed query:\n{fixed_query}")

            try:
                with connection.cursor() as cursor:
                    cursor.execute(fixed_query)
                    print(f"Fixed query for {i+1} executed successfully")

                    columns = [col[0] for col in cursor.description]
                    if "json_data" not in columns:
                        table_result = []
                    else:
                        idx = columns.index("json_data")
                        table_result = [row[idx] for row in cursor.fetchall()]

                    match = re.search(r'from\s+([a-zA-Z0-9_"]+)', fixed_query, re.IGNORECASE)
                    table_name = match.group(1).replace('"', '') if match else "unknown_table"

                    query_results.append({
                        "table": table_name,
                        "query": fixed_query,
                        "rows": table_result
                    })

            except Exception as e2:
                print(f"Fixed query still failed: {e2}")
                query_results.append({
                    "table": "error",
                    "query": fixed_query,
                    "error": str(e2)
                })
                continue

    
    print(query_results, "===query_results===")
    table_prompt = (
        "You are an expert at presenting data.\n\n"
        "Convert the following multiple table query results into HTML tables.\n"
        "Each table should include the table name as a header.\n"
        "If table is empty then skip, nothing to show.\n"
        "If there are errors, display them clearly under the table name.\n"
        "Always present the response strictly in a structured HTML table format.\n"
        "Return only runnable HTML table code, not including <pre><code>.\n"
        "Make sure to wrap the table in <div class='relative overflow-x-auto'></div> for responsiveness.\n"
        "Exclude the fields snapshot_at, loader_instance, database and table.\n"
        "Count the total rows found and shows at bottom. \n\n"
        f"Results:\n{query_results}"
    )

    html_table = api_engine(table_prompt, agent_data['online'], agent_data['llm'])
    return html_table


def process_excel_with_sequences(request, input_path: str, output_path: str, chat_id: str, default_value):
    wb = load_workbook(input_path)

    sequences_to_process = []

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str):
                    continue
                for match in GRAIC_SEQUENCE_PATTERN.finditer(cell.value):
                    seq_num = int(match.group(1))
                    new_chat_flag = match.group(2) or "N"
                    prompt_text = match.group(3).strip()
                    sequences_to_process.append({
                        "sheet": sheet,
                        "cell": cell,
                        "seq": seq_num,
                        "prompt": prompt_text,
                        "new_chat": new_chat_flag.upper(),
                        "full_match": match.group(0)
                    })

    sequences_to_process.sort(key=lambda x: x["seq"])

    previous_prompt = ""
    previous_response = ""

    for seq_info in sequences_to_process:
        sheet = seq_info["sheet"]
        cell = seq_info["cell"]
        row_index = cell.row
        col_index = cell.column
        col_letter = get_column_letter(col_index)
        cell_ref = f"{col_letter}{row_index}"

        prompt_text = seq_info["prompt"]
        new_chat = seq_info["new_chat"]
        full_match = seq_info["full_match"]

        print(f"🧠 Running sequence {seq_info['seq']} ({sheet.title}:{cell_ref})\nPrompt:\n{prompt_text}\n")

        agent = pick_agent_with_llm(prompt_text)
        print(f"→ Selected agent: {agent}")

        if agent == 'file_agent':
            generated_html = file_process(request, default_value, chat_id, agent,
                                          prompt=prompt_text, previous_response=previous_response,
                                          previous_prompt=previous_prompt, new_chat=new_chat)
        elif agent == 'sql_agent':
            generated_html = sql_process(request, default_value, chat_id, agent,
                                         prompt=prompt_text, previous_response=previous_response,
                                         previous_prompt=previous_prompt, new_chat=new_chat)
        else:
            generated_html = dt_process(request, default_value, chat_id, agent,
                                        prompt=prompt_text, previous_response=previous_response,
                                        previous_prompt=previous_prompt, new_chat=new_chat)

        previous_response = generated_html
        previous_prompt = prompt_text
        print(f"✅ Got response for sequence {seq_info['seq']}")

        is_table, parsed_data = parse_html_response(generated_html)

        if is_table:
            output_sheet_name = f"{sheet.title} ({cell_ref})"
            if len(output_sheet_name) > 31:
                output_sheet_name = output_sheet_name[:28] + "..."
            if output_sheet_name in wb.sheetnames:
                del wb[output_sheet_name]
            output_sheet = wb.create_sheet(output_sheet_name)

            output_sheet.append([f"Prompt from {sheet.title}!{cell_ref}"])
            output_sheet.append([prompt_text])
            output_sheet.append([])
            for table in parsed_data:
                for r in table:
                    output_sheet.append(r)
                output_sheet.append([])

            cell.value = cell.value.replace(full_match, f"See output sheet - {output_sheet_name}")

        else:
            cell.value = cell.value.replace(full_match, generated_html)

    wb.save(output_path)
    print(f"✅ File saved with processed sequences: {output_path}")


def ask_graic(request, chat_id):
    default_value = DefaultValues.objects.first()
    if not default_value:
        DefaultValues.objects.create()

    chat = get_object_or_404(Chat, uuid=chat_id, user=request.user)

    if request.method == 'POST':
        start_time = datetime.now()
        prompt = request.POST.get('prompt')
        uploaded_file = request.FILES.get('file')
        existing_file_id = request.POST.get('existing_file_id')

        generated_file_path = None
        summary_response = ""
        pre_response = ""
        
        print(uploaded_file, "===file===")

        agent = pick_agent_with_llm(prompt)
        print(agent, "===agent===")

        if uploaded_file:    
            with NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_in, \
                NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_out:

                for chunk in uploaded_file.chunks():
                    tmp_in.write(chunk)

                process_excel_with_sequences(request, tmp_in.name, tmp_out.name, chat.uuid, default_value)
                generated_file_path = tmp_out.name

        elif existing_file_id:
            file_instance = get_object_or_404(File, id=existing_file_id)
            with NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_in, \
                NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_out:

                with file_instance.file.open('rb') as f:
                    tmp_in.write(f.read())

                process_excel_with_sequences(request, tmp_in.name, tmp_out.name, chat.uuid, default_value)
                generated_file_path = tmp_out.name

        elif agent == 'file_agent':
            pre_response = file_process(request, default_value, chat.uuid, agent)
        elif agent == 'sql_agent':
            pre_response = sql_process(request, default_value, chat.uuid, agent)
        else:
            pre_response = dt_process(request, default_value, chat.uuid, agent)

        is_dt = False
        end_time = datetime.now()
        html_response = aggregate_responses(
            final_response=pre_response,
            prompt=prompt,
            user_id=request.user.id,
            chat_id=chat.uuid,
            metadata={
                'time_taken': (end_time - start_time).total_seconds(),
                'is_dt': is_dt,
                'summary_response': summary_response
            },
            generated_file_path=generated_file_path
        )

        return JsonResponse({
            "status": "complete",
            "response": html_response,
            "is_dt": False,
            'time_taken': (end_time - start_time).total_seconds()
        })

    return redirect(request.META.get('HTTP_REFERER'))

def save_preprompt(request):
    if request.method == 'POST':
        workflow = request.POST.get('workflow')
        pre_prompt = request.POST.get('pre_prompt')
        email_prompt = request.POST.get('email_prompt')

        PrePrompt.objects.update_or_create(
            workflow=workflow,
            defaults={
                'prompt': pre_prompt,
                'email_prompt': email_prompt
            }
        )

    return redirect(request.META.get('HTTP_REFERER'))

def get_prompt_by_workflow(request):
    workflow = request.GET.get('workflow')
    try:
        prompt_obj = PrePrompt.objects.get(workflow=workflow)
        return JsonResponse({'prompt': prompt_obj.prompt, 'email_prompt': prompt_obj.email_prompt})
    except PrePrompt.DoesNotExist:
        return JsonResponse({'error': 'Prompt not found.'}, status=404)


def graic_chat_delete(request, chat_id):
    if request.method == 'POST':
        chat = get_object_or_404(Chat, uuid=chat_id)
        chat.delete()
        return redirect(reverse('graic_chat_index'))
    
    return redirect(reverse('graic_chat_index'))


# Export

import io, os
import html2text
from docx import Document
from openpyxl import Workbook
from bs4 import BeautifulSoup
from django.http import HttpResponse, Http404
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
)
from reportlab.lib.units import cm
from django.core.files.base import ContentFile

def parse_html_response(html_text):
    if not html_text:
        return False, html_text

    soup = BeautifulSoup(html_text, "html.parser")
    tables = soup.find_all("table")

    if tables:
        all_tables = []
        for table in tables:
            rows = []
            for tr in table.find_all("tr"):
                cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                if cells:
                    rows.append(cells)
            all_tables.append(rows)
        return True, all_tables

    clean_text = soup.get_text(separator="\n", strip=True)
    return False, clean_text


def get_or_create_folder(name, parent=None, user=None):
    folder, _ = FileManager.objects.get_or_create(
        name=name,
        parent=parent,
        defaults={
            "created_by": user,
            "updated_by": user,
            "action_status": ActionStatus.IS_ACTIVE,
            "folder_status": DocumentStatus.APPROVED,
        }
    )
    return folder

def handle_export_destination(request, buffer, filename, content_type):
    destination = request.POST.get("destination", "download")
    buffer.seek(0)

    if destination != "grc":
        response = HttpResponse(buffer.getvalue(), content_type=content_type)
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    user = request.user
    username = user.username if user.is_authenticated else "anonymous"
    users_root = get_or_create_folder("users", parent=None, user=user)
    user_folder = get_or_create_folder(username, parent=users_root, user=user)
    graic_folder = get_or_create_folder("graic", parent=user_folder, user=user)

    file_obj = File.objects.create(
        file_manager=graic_folder,
        file=ContentFile(buffer.getvalue(), name=filename),
        action_status=ActionStatus.IS_ACTIVE,
        file_status=DocumentStatus.APPROVED,
        uploaded_by=user,
        updated_by=user
    )

    return JsonResponse({
        "success": True,
        "file_id": file_obj.id,
        "filename": filename,
        "path": f"users/{username}/graic"
    })


# ---- Export to Word ----
def export_graic_word(request, pk):
    entry = get_object_or_404(Graic, pk=pk)

    document = Document()
    document.add_heading("grAIc Chat Export", level=1)
    document.add_paragraph(f"You: {entry.prompt}", style='ListBullet')

    soup = BeautifulSoup(entry.response, "html.parser")
    tables = soup.find_all("table")

    for table_html in tables:
        rows = []
        for tr in table_html.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            if cells:
                rows.append(cells)

        if not rows:
            continue

        table = document.add_table(rows=0, cols=max(len(r) for r in rows))
        table.style = "Table Grid"

        for row_data in rows:
            cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                cells[i].text = cell_text

    buffer = io.BytesIO()
    document.save(buffer)

    return handle_export_destination(
        request,
        buffer,
        f"graic_{entry.id}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ---- Export to Excel ----
def export_graic_excel(request, pk):
    entry = get_object_or_404(Graic, pk=pk)
    is_table, parsed_data = parse_html_response(entry.response)

    wb = Workbook()
    ws = wb.active
    ws.title = "grAIc Export"

    if is_table:
        for table in parsed_data:
            for row in table:
                ws.append(row)
            ws.append([])
    else:
        ws.append(["You:", entry.prompt])
        ws.append(["grAIc:", parsed_data])

    buffer = io.BytesIO()
    wb.save(buffer)

    return handle_export_destination(
        request,
        buffer,
        f"graic_{entry.id}.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


def export_generated_file(request, pk):
    entry = get_object_or_404(Graic, pk=pk)

    if not entry.generated_file:
        raise Http404("No generated file found")

    destination = request.POST.get("destination", "download")
    source_file = entry.generated_file

    filename = os.path.basename(source_file.name)
    content_type = "application/octet-stream"

    source_file.open("rb")
    file_bytes = source_file.read()
    source_file.close()

    if destination != "grc":
        response = HttpResponse(file_bytes, content_type=content_type)
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    user = request.user
    username = user.username if user.is_authenticated else "anonymous"

    users_root = get_or_create_folder("users", parent=None, user=user)
    user_folder = get_or_create_folder(username, parent=users_root, user=user)
    graic_folder = get_or_create_folder("graic", parent=user_folder, user=user)

    file_obj = File.objects.create(
        file_manager=graic_folder,
        file=ContentFile(file_bytes, name=filename),
        action_status=ActionStatus.IS_ACTIVE,
        file_status=DocumentStatus.APPROVED,
        uploaded_by=user,
        updated_by=user,
    )

    return JsonResponse({
        "success": True,
        "file_id": file_obj.id,
        "filename": filename,
        "path": f"users/{username}/graic"
    })


# ---- Export to PDF ----
def export_graic_pdf(request, pk):
    entry = get_object_or_404(Graic, pk=pk)
    soup = BeautifulSoup(entry.response, "html.parser")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=1*cm,
        rightMargin=1*cm,
        topMargin=1*cm,
        bottomMargin=1*cm,
    )

    styles = getSampleStyleSheet()
    normal_style = styles["Normal"]
    title_style = ParagraphStyle(
        "TableTitle",
        parent=styles["Heading2"],
        alignment=1,
        fontSize=12,
        spaceAfter=6,
    )

    story = []

    for table in soup.find_all("table"):
        headers, rows = [], []

        thead = table.find("thead")
        if thead:
            for tr in thead.find_all("tr"):
                headers.append([
                    Paragraph(td.get_text(strip=True), normal_style)
                    for td in tr.find_all(["th", "td"])
                ])

        tbody = table.find("tbody")
        if tbody:
            for tr in tbody.find_all("tr"):
                rows.append([
                    Paragraph(td.get_text(strip=True) or "", normal_style)
                    for td in tr.find_all(["td", "th"])
                ])

        data = headers + rows
        if not data:
            continue

        col_width = doc.width / len(data[0])
        table_obj = Table(data, repeatRows=1, colWidths=[col_width]*len(data[0]))
        table_obj.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
        ]))

        story.append(table_obj)
        story.append(Spacer(1, 12))

    doc.build(story)

    return handle_export_destination(
        request,
        buffer,
        f"graic_{pk}.pdf",
        "application/pdf"
    )


# ---- Export to Markdown ----
def export_graic_markdown(request, pk):
    entry = get_object_or_404(Graic, pk=pk)

    converter = html2text.HTML2Text()
    converter.ignore_links = False
    converter.body_width = 0

    md_content = f"# grAIc Chat Export\n\n"
    md_content += f"**You:** {entry.prompt}\n\n"

    is_table, parsed_data = parse_html_response(entry.response)

    if is_table:
        for idx, table in enumerate(parsed_data, 1):
            if not table:
                continue

            header = table[0]
            md_content += f"**Table {idx}:**\n\n"
            md_content += "| " + " | ".join(header) + " |\n"
            md_content += "| " + " | ".join(["---"] * len(header)) + " |\n"

            for row in table[1:]:
                row += [""] * (len(header) - len(row))
                md_content += "| " + " | ".join(row) + " |\n"

            md_content += "\n"
    else:
        md_content += converter.handle(parsed_data or "")

    buffer = io.BytesIO(md_content.encode("utf-8"))

    return handle_export_destination(
        request,
        buffer,
        f"graic_{entry.id}.md",
        "text/markdown"
    )
